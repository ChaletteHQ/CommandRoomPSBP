#!/usr/bin/env python3
"""
Tests for `shared/scripts/brief_writer.py` (v2.14.32+).

Covers the contract `brief_writer` must satisfy:

  - make_brief() writes a non-empty .docx to the requested path
  - Eyebrow label is "CALL PREP" for call_prep, "MEETING BRIEF" for past_meeting
  - Title appears as a 22pt bold paragraph in dark navy
  - Section headings appear as 13pt bold paragraphs in dark navy
  - Body paragraphs render in 11pt near-black with 1.25 line spacing
  - Bullets use the "List Bullet" style
  - Footer is hard-coded to "Command Room" (centered)
  - Margins are 0.9" top/bottom, 1.0" left/right
  - JSON-via-stdin entry point (`make_brief_from_json`) round-trips
  - Bad inputs raise ValueError (kind, missing fields)

Forwardable-clean structural enforcement (per CONTRACT.md Rule 15):
  - Footer never accepts provenance metadata — it's a fixed positional argument
  - The pre-v2.14.32 `Source: ... | Fired: ... | Inputs: ... | TTL: ...`
    pattern cannot leak through brief_writer.
"""
import json
import os
import sys
import tempfile
import traceback

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(HERE, "..", "shared", "scripts"))

from brief_writer import make_brief, make_brief_from_json
from brand import DEFAULT_BRAND
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


results = {"pass": 0, "fail": 0, "failures": []}


def check(name, condition, detail=""):
    if condition:
        results["pass"] += 1
        print(f"  PASS  {name}")
    else:
        results["fail"] += 1
        results["failures"].append(f"{name} ({detail})" if detail else name)
        print(f"  FAIL  {name}{(' — ' + detail) if detail else ''}")


def _sample_call_prep(path):
    return make_brief(
        path,
        brief_kind="call_prep",
        title="Sam Sample — Q2 deck review",
        subtitle="Friday, May 9, 2026 · 9:00 AM PT · Summit Company",
        # OUT2 §4 flip: call_prep is a STANDARD_KIND — exec_header required.
        exec_header={"verdict": "Walk out with the Q2 deck approved."},
        sections=[
            {"heading": "Lead with", "body": "Open warm.\n\nPivot to substance fast."},
            {"heading": "What to demo", "bullets": ["Connect Gmail.", "Surface a real signal."]},
            {"heading": "Suggested outcome", "body": "End with an installed Cowork."},
        ],
        # B3: this fixture exercises typography/layout, not the output contract;
        # the contract gate is covered by run_output_contract_validator_test.py.
        contract="off",
    )


def _sample_past_meeting(path):
    return make_brief(
        path,
        brief_kind="past_meeting",
        title="Sam Sample — UX review (continuation)",
        subtitle="Wednesday, Apr 29, 2026 · 9:05 PM PT · 70 min · Command Room",
        sections=[
            {"heading": "Attendees", "bullets": ["Matthew Sample", "Sam Sample"]},
            {"heading": "Summary", "bullets": ["Locked the 5-promise framing.", "Sam ran Inbox end-to-end."]},
            {"heading": "Decisions", "bullets": ["Promise #3 wording: 'nothing important slips'."]},
        ],
    )


# ============================================================================
# Test 1 — basic file output
# ============================================================================
print("\n=== make_brief: basic output ===")

with tempfile.TemporaryDirectory() as tmp:
    cp = os.path.join(tmp, "cp.docx")
    pm = os.path.join(tmp, "pm.docx")
    _sample_call_prep(cp)
    _sample_past_meeting(pm)
    check("call_prep file written", os.path.isfile(cp))
    check("past_meeting file written", os.path.isfile(pm))
    check("call_prep file is non-empty", os.path.getsize(cp) > 5000)
    check("past_meeting file is non-empty", os.path.getsize(pm) > 5000)

    # ============================================================================
    # Test 2 — eyebrow label correctness
    # ============================================================================
    print("\n=== eyebrow labels ===")
    doc_cp = Document(cp)
    doc_pm = Document(pm)
    eyebrow_cp = doc_cp.paragraphs[0].text
    eyebrow_pm = doc_pm.paragraphs[0].text
    check("call_prep eyebrow = 'CALL PREP'", eyebrow_cp == "CALL PREP", f"got {eyebrow_cp!r}")
    check("past_meeting eyebrow = 'MEETING BRIEF'", eyebrow_pm == "MEETING BRIEF", f"got {eyebrow_pm!r}")

    # ============================================================================
    # Test 3 — title typography (22pt bold, dark navy)
    # ============================================================================
    print("\n=== title typography ===")
    title_para = doc_cp.paragraphs[1]
    title_run = title_para.runs[0]
    check("title text matches", title_para.text == "Sam Sample — Q2 deck review")
    check("title is 22pt", title_run.font.size == Pt(22), f"got {title_run.font.size}")
    check("title is bold", title_run.font.bold is True)
    # SPEC OUT1 deliberate default-theme diff: heading navy refined
    # 0F2A3F -> 102A40 (DEFAULT_BRAND.palette.heading). Anchored to the brand
    # source of truth so a future default-theme change updates in one place.
    check(
        "title color is the default-brand heading navy",
        str(title_run.font.color.rgb) == DEFAULT_BRAND["palette"]["heading"],
        f"got {title_run.font.color.rgb}, expected {DEFAULT_BRAND['palette']['heading']}",
    )

    # ============================================================================
    # Test 4 — subtitle typography (11pt muted grey)
    # ============================================================================
    print("\n=== subtitle typography ===")
    sub_run = doc_cp.paragraphs[2].runs[0]
    check("subtitle is 11pt", sub_run.font.size == Pt(11))
    check(
        "subtitle color is muted grey (#6B6B6B)",
        str(sub_run.font.color.rgb) == "6B6B6B",
    )

    # ============================================================================
    # Test 5 — section heading typography (13pt bold dark navy)
    # ============================================================================
    print("\n=== section heading typography ===")
    # Find the first section heading paragraph (after eyebrow, title, subtitle, rule)
    heading_paras = [
        p for p in doc_cp.paragraphs
        if p.runs and p.runs[0].font.size == Pt(13) and p.runs[0].font.bold
    ]
    check("at least one section heading present", len(heading_paras) >= 3)
    if heading_paras:
        h = heading_paras[0]
        check("first section heading text = 'Lead with'", h.text == "Lead with")
        check(
            "section heading color is the default-brand heading navy",
            str(h.runs[0].font.color.rgb) == DEFAULT_BRAND["palette"]["heading"],
        )

    # ============================================================================
    # Test 6 — bullets use List Bullet style
    # ============================================================================
    print("\n=== bullets ===")
    bullet_paras = [p for p in doc_cp.paragraphs if p.style.name == "List Bullet"]
    check("at least 2 List Bullet paragraphs", len(bullet_paras) >= 2)
    if bullet_paras:
        check(
            "first bullet text matches",
            bullet_paras[0].text == "Connect Gmail.",
            f"got {bullet_paras[0].text!r}",
        )

    # ============================================================================
    # Test 7 — footer is hard-coded "Command Room" centered
    # ============================================================================
    print("\n=== footer enforcement ===")
    footer_para = doc_cp.sections[0].footer.paragraphs[0]
    check("footer text = 'Command Room'", footer_para.text == "Command Room")
    check("footer is centered", footer_para.alignment == WD_ALIGN_PARAGRAPH.CENTER)
    footer_run = footer_para.runs[0]
    check("footer color is muted grey", str(footer_run.font.color.rgb) == "6B6B6B")
    check("footer is 9pt", footer_run.font.size == Pt(9))

    # ============================================================================
    # Test 8 — margins (0.9" top/bot, 1.0" left/right)
    # ============================================================================
    print("\n=== margins ===")
    sec = doc_cp.sections[0]
    # 1 inch = 914400 EMU; 0.9 inch = 822960 EMU
    check("top margin is 0.9 inch", sec.top_margin == 822960, f"got {sec.top_margin}")
    check("bottom margin is 0.9 inch", sec.bottom_margin == 822960)
    check("left margin is 1.0 inch", sec.left_margin == 914400)
    check("right margin is 1.0 inch", sec.right_margin == 914400)

    # ============================================================================
    # Test 9 — forwardable-clean structural enforcement
    # ============================================================================
    print("\n=== forwardable-clean enforcement ===")
    # Verify no provenance leak strings appear anywhere in the document body
    full_text = "\n".join(p.text for p in doc_cp.paragraphs)
    check("no 'Source:' leak in body", "Source:" not in full_text)
    check("no 'Fired:' leak in body", "Fired:" not in full_text)
    check("no 'TTL:' leak in body", "TTL:" not in full_text)
    check("no 'Inputs:' leak in body", "Inputs:" not in full_text)
    check("no 'meeting_id:' leak in body", "meeting_id:" not in full_text)

    # ============================================================================
    # Test 10 — JSON-via-stdin entry point
    # ============================================================================
    print("\n=== make_brief_from_json ===")
    json_path = os.path.join(tmp, "from_json.docx")
    payload = json.dumps({
        "output_path": json_path,
        "brief_kind": "call_prep",
        "title": "JSON Test Title",
        "subtitle": "Today",
        "exec_header": {"verdict": "Walk out with the JSON path proven."},  # OUT2 §4 flip
        "sections": [{"heading": "Solo section", "body": "Hello."}],
        "contract": "off",  # B3: exercises the JSON round-trip, not the contract gate
    })
    returned = make_brief_from_json(payload)
    check("JSON entry point returns the path", returned == json_path)
    check("JSON-produced file exists", os.path.isfile(json_path))
    doc_json = Document(json_path)
    check("JSON-produced file has the right title", doc_json.paragraphs[1].text == "JSON Test Title")


# ============================================================================
# Test 11 — bad-input rejection
# ============================================================================
print("\n=== input validation ===")

with tempfile.TemporaryDirectory() as tmp:
    bad = os.path.join(tmp, "bad.docx")

    def _expect_value_error(label, fn):
        try:
            fn()
            check(label, False, "expected ValueError, got success")
        except ValueError:
            check(label, True)
        except Exception as e:
            check(label, False, f"expected ValueError, got {type(e).__name__}: {e}")

    _expect_value_error(
        "rejects unknown brief_kind",
        lambda: make_brief(bad, brief_kind="totally_made_up_kind", title="t", subtitle="s",
                           sections=[{"heading": "h", "body": "b"}]),
    )
    _expect_value_error(
        "rejects missing title",
        lambda: make_brief(bad, brief_kind="call_prep", title="", subtitle="s",
                           sections=[{"heading": "h", "body": "b"}]),
    )
    _expect_value_error(
        "rejects empty sections",
        lambda: make_brief(bad, brief_kind="call_prep", title="t", subtitle="s", sections=[]),
    )
    _expect_value_error(
        "rejects section without heading",
        lambda: make_brief(bad, brief_kind="call_prep", title="t", subtitle="s",
                           sections=[{"body": "b"}], contract="off"),
    )
    _expect_value_error(
        "rejects section without body or bullets",
        lambda: make_brief(bad, brief_kind="call_prep", title="t", subtitle="s",
                           sections=[{"heading": "h"}], contract="off"),
    )


# ============================================================================
# Test 12 — SPEC EXEC1: exec_header renders before the first section
# ============================================================================
print("\n=== EXEC1: exec_header rendering ===")

with tempfile.TemporaryDirectory() as tmp:
    p = os.path.join(tmp, "exec.docx")
    # FS-13: the CHANGED/DECIDE/NEEDED eyebrow is brief-family only. weekly_recap
    # is a brief-family STANDARD_KIND, so it renders the full eyebrow (memo /
    # one_pager render the verdict lead ONLY — see the dedicated per-kind test).
    make_brief(
        p,
        brief_kind="weekly_recap",
        title="Acme renewal is the only call this week",
        subtitle="Sat, Jun 14, 2026",
        exec_header={
            "verdict": "Ratify the Acme renewal by Friday.",
            "changed": "Acme moved from handshake to paperwork.",
            "decide": "Whether to gate the sales hire — by Jun 15.",
            "needs": "Approve the redline below.",
        },
        sections=[
            {"heading": "Recommendation", "body": "Ratify the renewal."},
            {"heading": "Why", "body": "Because the terms hold."},
        ],
        contract="off",
        voice_gate="off",
    )
    doc = Document(p)
    texts = [pp.text for pp in doc.paragraphs if pp.text.strip()]
    # eyebrow, title, subtitle, verdict, CHANGED, DECIDE, NEEDED, then sections
    check("verdict line present", "Ratify the Acme renewal by Friday." in texts)
    check("CHANGED line present + labeled",
          any(t.startswith("CHANGED") and "paperwork" in t for t in texts),
          f"texts={texts[:8]}")
    check("DECIDE line present + labeled", any(t.startswith("DECIDE") for t in texts))
    check("NEEDED line present + labeled", any(t.startswith("NEEDED") for t in texts))
    # exec header must precede the first section heading
    verdict_idx = texts.index("Ratify the Acme renewal by Friday.")
    rec_idx = texts.index("Recommendation")
    check("exec header precedes first section", verdict_idx < rec_idx)
    # verdict run is bold
    verdict_para = [pp for pp in doc.paragraphs if pp.text == "Ratify the Acme renewal by Friday."][0]
    check("verdict run is bold", verdict_para.runs[0].font.bold is True)

# ============================================================================
# Test 13 — SPEC EXEC1: asks block renders last under canonical heading
# ============================================================================
print("\n=== EXEC1: asks block ===")

with tempfile.TemporaryDirectory() as tmp:
    p = os.path.join(tmp, "asks.docx")
    make_brief(
        p,
        brief_kind="memo",
        title="t",
        subtitle="s",
        exec_header={"verdict": "Do the thing."},
        sections=[{"heading": "Body", "body": "content"}],
        asks=[{"text": "Approve the redline", "deadline": "Friday"},
              {"text": "Confirm the board date"}],
        contract="off",
        voice_gate="off",
    )
    doc = Document(p)
    texts = [pp.text for pp in doc.paragraphs if pp.text.strip()]
    check("ASK heading present", "What I need from you" in texts)
    bullet_texts = [pp.text for pp in doc.paragraphs if pp.style.name == "List Bullet"]
    check("ask with deadline renders ' — by Friday'",
          any("Approve the redline — by Friday" == b for b in bullet_texts),
          f"bullets={bullet_texts}")
    check("ask without deadline renders plain",
          any(b == "Confirm the board date" for b in bullet_texts))
    # asks block comes after the body section
    check("asks heading is after body section",
          texts.index("What I need from you") > texts.index("Body"))

    # Zero / empty asks → no ASK heading rendered
    p2 = os.path.join(tmp, "noasks.docx")
    make_brief(p2, brief_kind="memo", title="t", subtitle="s",
               exec_header={"verdict": "v"},
               sections=[{"heading": "Body", "body": "c"}], asks=[],
               contract="off", voice_gate="off")
    texts2 = [pp.text for pp in Document(p2).paragraphs]
    check("zero asks → no ASK heading", "What I need from you" not in texts2)


# ============================================================================
# Test 14 — SPEC EXEC1: asks cap (>3) raises
# ============================================================================
print("\n=== EXEC1: asks cap ===")

with tempfile.TemporaryDirectory() as tmp:
    bad = os.path.join(tmp, "bad.docx")

    def _expect_ve(label, fn):
        try:
            fn()
            check(label, False, "expected ValueError")
        except ValueError:
            check(label, True)
        except Exception as e:
            check(label, False, f"expected ValueError, got {type(e).__name__}: {e}")

    _expect_ve(
        "asks >3 raises",
        lambda: make_brief(bad, brief_kind="memo", title="t", subtitle="s",
                           exec_header={"verdict": "v"},  # so the raise is the asks cap
                           sections=[{"heading": "h", "body": "b"}],
                           asks=[{"text": "a"}, {"text": "b"}, {"text": "c"}, {"text": "d"}],
                           contract="off", voice_gate="off"),
    )
    _expect_ve(
        "ask without text raises",
        lambda: make_brief(bad, brief_kind="memo", title="t", subtitle="s",
                           exec_header={"verdict": "v"},  # so the raise is the asks shape
                           sections=[{"heading": "h", "body": "b"}],
                           asks=[{"deadline": "Friday"}], contract="off", voice_gate="off"),
    )
    # exactly 3 asks is OK
    ok = os.path.join(tmp, "ok.docx")
    make_brief(ok, brief_kind="memo", title="t", subtitle="s",
               exec_header={"verdict": "v"},
               sections=[{"heading": "h", "body": "b"}],
               asks=[{"text": "a"}, {"text": "b"}, {"text": "c"}],
               contract="off", voice_gate="off")
    check("exactly 3 asks is allowed", os.path.isfile(ok))


# ============================================================================
# Test 15 — SPEC EXEC1: recommendation-ordering check (decision-shaped kinds)
# ============================================================================
print("\n=== EXEC1: recommendation ordering ===")

with tempfile.TemporaryDirectory() as tmp:
    bad = os.path.join(tmp, "bad.docx")

    def _expect_ve2(label, fn):
        try:
            fn()
            check(label, False, "expected ValueError")
        except ValueError:
            check(label, True)
        except Exception as e:
            check(label, False, f"expected ValueError, got {type(e).__name__}: {e}")

    _expect_ve2(
        "decision_memo with late Recommendation (idx 4) raises",
        lambda: make_brief(bad, brief_kind="decision_memo", title="t", subtitle="s",
                           exec_header={"verdict": "v"},  # so the raise is the ordering check
                           sections=[{"heading": "Framing", "body": "b"},
                                     {"heading": "Options", "body": "b"},
                                     {"heading": "Criteria", "body": "b"},
                                     {"heading": "Comparison", "body": "b"},
                                     {"heading": "Recommendation", "body": "b"}],
                           contract="off", voice_gate="off"),
    )
    # rec in first three sections → OK
    ok = os.path.join(tmp, "ok.docx")
    make_brief(ok, brief_kind="decision_memo", title="t", subtitle="s",
               exec_header={"verdict": "v"},
               sections=[{"heading": "Recommendation", "body": "b"},
                         {"heading": "Comparison", "body": "b"}],
               contract="off", voice_gate="off")
    check("decision_memo with early Recommendation passes", os.path.isfile(ok))

    # non-decision kind (call_prep) is NOT subject to the ordering check
    ok2 = os.path.join(tmp, "ok2.docx")
    make_brief(ok2, brief_kind="call_prep", title="t", subtitle="s",
               exec_header={"verdict": "v"},  # OUT2 §4 flip: call_prep requires it
               sections=[{"heading": "A", "body": "b"}, {"heading": "B", "body": "b"},
                         {"heading": "C", "body": "b"}, {"heading": "Decisions", "body": "b"}],
               contract="off", voice_gate="off")
    check("call_prep late 'Decisions' section is NOT ordering-checked", os.path.isfile(ok2))


# ============================================================================
# Test 16 — SPEC OUT2 §4: STANDARD_KINDS exec_header is REQUIRED (the deferred
#           EXEC1 release-N+1 flip) — ValueError naming the kind, no partial
#           file, brief_meta severity="error" audit when workspace_root passed
# ============================================================================
print("\n=== OUT2 §4: exec_header hard-require (the flip) ===")

from brief_writer import STANDARD_KINDS  # noqa: E402

with tempfile.TemporaryDirectory() as tmp:
    # EVERY kind in the set must refuse to render without an exec_header, and
    # the error must NAME the kind (spec §6: "or the test fails naming the
    # skill"). No partial file may exist after the raise.
    for kind in sorted(STANDARD_KINDS):
        p = os.path.join(tmp, f"noheader_{kind}.docx")
        try:
            make_brief(p, brief_kind=kind, title="t", subtitle="s",
                       sections=[{"heading": "h", "body": "b"}],
                       contract="off", voice_gate="off")
            check(f"STANDARD_KIND {kind} without exec_header raises", False)
        except ValueError as e:
            check(f"STANDARD_KIND {kind} without exec_header raises naming the kind",
                  kind in str(e) and "exec_header" in str(e))
        check(f"STANDARD_KIND {kind} refusal leaves no partial file",
              not os.path.isfile(p))

    # An empty/whitespace verdict is the same as no header.
    p_blank = os.path.join(tmp, "blankverdict.docx")
    try:
        make_brief(p_blank, brief_kind="memo", title="t", subtitle="s",
                   exec_header={"verdict": "   "},
                   sections=[{"heading": "h", "body": "b"}],
                   contract="off", voice_gate="off")
        check("blank verdict raises like a missing header", False)
    except ValueError:
        check("blank verdict raises like a missing header", True)

    # NON-standard kinds still render freely without one (never required).
    p_free = os.path.join(tmp, "free.docx")
    make_brief(p_free, brief_kind="past_meeting", title="t", subtitle="s",
               sections=[{"heading": "h", "body": "b"}],
               contract="off", voice_gate="off")
    check("non-STANDARD kind renders without exec_header", os.path.isfile(p_free))

    # When workspace_root is provided, the refusal leaves a brief_meta
    # severity="error" audit event — the substrate trace of a lagging caller.
    ws = os.path.join(tmp, "ws")
    os.makedirs(os.path.join(ws, "_hq", "data"), exist_ok=True)
    p2 = os.path.join(tmp, "audit.docx")
    try:
        make_brief(p2, brief_kind="board_pack", title="t", subtitle="s",
                   sections=[{"heading": "h", "body": "b"}],
                   contract="off", voice_gate="off", workspace_root=ws)
    except ValueError:
        pass
    events_path = os.path.join(ws, "_hq", "data", "events.jsonl")
    check("brief_meta audit event file written on refusal", os.path.isfile(events_path))
    if os.path.isfile(events_path):
        lines = [l for l in open(events_path, encoding="utf-8").read().splitlines() if l.strip()]
        evs = [json.loads(l) for l in lines]
        bm = [e for e in evs if e.get("type") == "brief_meta"]
        check("brief_meta carries the kind + error severity",
              bool(bm) and bm[0]["data"].get("brief_kind") == "board_pack"
              and bm[0]["data"].get("severity") == "error")

    # A STANDARD_KIND WITH exec_header saves cleanly, NO brief_meta event.
    ws2 = os.path.join(tmp, "ws2")
    os.makedirs(os.path.join(ws2, "_hq", "data"), exist_ok=True)
    p3 = os.path.join(tmp, "withheader.docx")
    make_brief(p3, brief_kind="memo", title="t", subtitle="s",
               exec_header={"verdict": "Ship it."},
               sections=[{"heading": "h", "body": "b"}],
               contract="off", voice_gate="off", workspace_root=ws2)
    check("exec_header present → saves", os.path.isfile(p3))
    ep2 = os.path.join(ws2, "_hq", "data", "events.jsonl")
    no_warn = (not os.path.isfile(ep2)) or all(
        json.loads(l).get("type") != "brief_meta"
        for l in open(ep2, encoding="utf-8").read().splitlines() if l.strip()
    )
    check("exec_header present → no brief_meta event", no_warn)


# ============================================================================
# Test 17 — SPEC C1: value_receipt kind is accepted + eyebrow "VALUE RECEIPT"
# ============================================================================
print("\n=== C1: value_receipt brief_kind ===")

with tempfile.TemporaryDirectory() as tmp:
    vr = os.path.join(tmp, "vr.docx")
    make_brief(
        vr,
        brief_kind="value_receipt",
        title="Value Receipt — May 2026",
        subtitle="Your operating layer, in numbers",
        sections=[
            {"heading": "What Command Room handled",
             "bullets": ["14 commitments captured that weren't tracked anywhere else",
                         "12 meetings turned into a structured brief"]},
            {"heading": "Time absorbed",
             "body": "~31 hours of operational overhead absorbed.\n\nConservative — assumes you would have done each of these tasks yourself at average speed."},
        ],
    )
    check("value_receipt file written", os.path.isfile(vr))
    doc_vr = Document(vr)
    check("value_receipt eyebrow = 'VALUE RECEIPT'",
          doc_vr.paragraphs[0].text == "VALUE RECEIPT", f"got {doc_vr.paragraphs[0].text!r}")
    body_text = "\n".join(p.text for p in doc_vr.paragraphs)
    check("forwardable doc contains 'Conservative'", "Conservative" in body_text)


# ============================================================================
# Summary
# ============================================================================
print(f"\n=== {results['pass']} passed, {results['fail']} failed ===")
if results["fail"]:
    print("Failures:")
    for f in results["failures"]:
        print(f"  - {f}")
    sys.exit(1)
sys.exit(0)
