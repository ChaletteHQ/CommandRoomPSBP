#!/usr/bin/env python3
"""SPEC OUT2 §5 — the cross-skill output profile.

Covers (house convention: check(name, cond) prints OK/FAIL, exit 1 on any
failure, auto-discovered by run_all.py):

  - get_output_profile() defaults: no workspace / absent file / malformed file
    all resolve to DEFAULT_OUTPUT_PROFILE, byte-stably, as a fresh copy
  - the skill_config_writer wrapper shape AND a raw dict both resolve
  - invalid values (typo density, unknown format, bad page cap) keep defaults
  - validate_output_profile flags unknown keys / bad values, passes clean partials
  - BYTE-STABILITY GOLDEN: a no-profile render == an explicit default-profile
    render (structure + spacing signature) — absent file = today's behavior
  - density "narrative" loosens body spacing; defaults restored after (no bleed)
  - visual_bias "prose_first" flips tiles/body order within a section
  - page_cap: over-cap render warns on stderr and STILL SAVES (warn-only forever)
  - schema registration: output_profile + the four OUT2 §5 composers are in
    skill_config.schema.json; save_skill_config round-trips a profile write and
    rejects an unknown profile key
  - ⛔ FENCE: no onboarding surface mentions the output profile, and
    FIRST_RUN_PROTOCOL's catalog explicitly excludes it
"""
from __future__ import annotations

import io
import json
import sys
import tempfile
from contextlib import redirect_stderr
from pathlib import Path

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
sys.path.insert(0, str(ROOT / "shared" / "scripts"))

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

import brief_writer as bw  # noqa: E402
import skill_config_writer as scw  # noqa: E402
from output_profile import (  # noqa: E402
    DEFAULT_OUTPUT_PROFILE,
    get_output_profile,
    validate_output_profile,
)
from docx import Document  # noqa: E402

_failures: list = []


def check(name: str, cond: bool) -> None:
    print(f"  {'OK  ' if cond else 'FAIL'} {name}")
    if not cond:
        _failures.append(name)


def _ws() -> Path:
    ws = Path(tempfile.mkdtemp(prefix="outprof_ws_"))
    (ws / "_hq" / "data" / "skill_config").mkdir(parents=True)
    (ws / "_hq" / "data" / "events.jsonl").write_text("", encoding="utf-8")
    return ws


def _write_profile(ws: Path, cfg: dict, wrapped: bool = True) -> None:
    path = ws / "_hq" / "data" / "skill_config" / "output_profile.json"
    payload = {"schema_version": 1, "skill_name": "output_profile",
               "config": cfg} if wrapped else cfg
    path.write_text(json.dumps(payload), encoding="utf-8")


def _sections():
    return [
        {"heading": "Summary",
         "tiles": [{"label": "Handled", "value": "12"}],
         "body": "First body paragraph with several words in it.\n\nSecond paragraph."},
        {"heading": "Detail", "body": "Detail body text for the spacing check."},
    ]


def _render(path: str, ws: Path | None) -> str:
    bw.make_brief(path, brief_kind="memo", title="Profile Test", subtitle="sub",
                  exec_header={"verdict": "Profile check."},
                  sections=_sections(), contract="off", voice_gate="off",
                  workspace_root=str(ws) if ws else None)
    return path


_BODY_TEXT_STEMS = ("First body paragraph", "Second paragraph", "Detail body")


def _spacing_signature(path: str) -> list:
    """(text, line_spacing, space_after_pt) for the section BODY paragraphs
    (identified by their known test texts — the exec header also renders at
    11pt but carries its own fixed spacing, which the profile never touches)."""
    doc = Document(path)
    sig = []
    for p in doc.paragraphs:
        if p.text.startswith(_BODY_TEXT_STEMS):
            sa = p.paragraph_format.space_after
            sig.append((p.text, p.paragraph_format.line_spacing,
                        int(sa.pt) if sa is not None else None))
    return sig


def _block_order(path: str) -> list:
    """Document body child order: 'p:<text>' for paragraphs, 'tbl' for tables —
    enough to see whether tiles (a table) precede or follow the body text."""
    doc = Document(path)
    order = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "tbl":
            order.append("tbl")
        elif tag == "p":
            text = "".join(node.text or "" for node in child.iter()
                           if node.tag.endswith("}t"))
            order.append(f"p:{text}")
    return order


def main() -> int:
    print("=== SPEC OUT2 §5: output profile ===")

    # ---- resolution defaults ----
    d = get_output_profile(None)
    check("no workspace -> DEFAULT_OUTPUT_PROFILE", d == DEFAULT_OUTPUT_PROFILE)
    d["density"] = "MUTATED"
    check("returns a fresh copy (no shared mutation)",
          DEFAULT_OUTPUT_PROFILE["density"] == "tight")
    ws = _ws()
    check("absent file -> defaults", get_output_profile(ws) == DEFAULT_OUTPUT_PROFILE)
    (ws / "_hq" / "data" / "skill_config" / "output_profile.json").write_text(
        "{not json", encoding="utf-8")
    check("malformed file -> defaults (never raises)",
          get_output_profile(ws) == DEFAULT_OUTPUT_PROFILE)

    # ---- wrapper + raw shapes, partial override, typo safety ----
    ws = _ws()
    _write_profile(ws, {"density": "narrative"}, wrapped=True)
    got = get_output_profile(ws)
    check("wrapper shape: saved key wins", got["density"] == "narrative")
    check("wrapper shape: unset keys keep defaults",
          got["visual_bias"] == "tiles_first" and got["default_format"] == "docx")
    _write_profile(ws, {"visual_bias": "prose_first"}, wrapped=False)
    check("raw-dict shape also resolves",
          get_output_profile(ws)["visual_bias"] == "prose_first")
    _write_profile(ws, {"density": "airy", "default_format": "pdf",
                        "page_cap": {"memo": 0, "board_pack": 8, "x": "two"},
                        "unknown_knob": True})
    got = get_output_profile(ws)
    check("invalid density value -> default kept", got["density"] == "tight")
    check("unknown format -> docx kept", got["default_format"] == "docx")
    check("page_cap: only positive-int entries survive",
          got["page_cap"] == {"board_pack": 8})
    check("unknown key ignored at resolution", "unknown_knob" not in got)

    # ---- validate_output_profile ----
    check("validate: full default clean", validate_output_profile(DEFAULT_OUTPUT_PROFILE) == [])
    check("validate: clean partial passes",
          validate_output_profile({"density": "narrative"}) == [])
    check("validate: bad density flagged",
          any("density" in p for p in validate_output_profile({"density": "airy"})))
    check("validate: unknown key flagged",
          any("nope" in p for p in validate_output_profile({"nope": 1})))
    check("validate: non-docx format flagged (Wave 3 fence)",
          any("default_format" in p for p in validate_output_profile({"default_format": "html"})))
    check("validate: bad page_cap flagged",
          any("page_cap" in p for p in validate_output_profile({"page_cap": {"memo": 0}})))
    check("validate: non-dict rejected", validate_output_profile("x") != [])

    outdir = Path(tempfile.mkdtemp(prefix="outprof_out_"))

    # ---- GOLDEN byte-stability: absent profile == explicit defaults == today ----
    p_none = _render(str(outdir / "none.docx"), None)
    ws = _ws()  # no profile file
    p_absent = _render(str(outdir / "absent.docx"), ws)
    _write_profile(ws, dict(DEFAULT_OUTPUT_PROFILE))
    p_expl = _render(str(outdir / "explicit_defaults.docx"), ws)
    sig_none = _spacing_signature(p_none)
    check("golden: absent-profile render == no-workspace render (spacing+text)",
          _spacing_signature(p_absent) == sig_none)
    check("golden: explicit-default-profile render == no-workspace render",
          _spacing_signature(p_expl) == sig_none)
    check("golden: block order identical across the three",
          _block_order(p_none) == _block_order(p_absent) == _block_order(p_expl))
    check("golden: default body spacing is the pre-profile 1.25 / 6pt",
          all(ls == 1.25 and sa == 6 for _, ls, sa in sig_none))

    # ---- density: narrative loosens body spacing; no bleed after ----
    ws = _ws()
    _write_profile(ws, {"density": "narrative"})
    p_narr = _render(str(outdir / "narrative.docx"), ws)
    sig_narr = _spacing_signature(p_narr)
    check("narrative: body line spacing loosened (1.40)",
          all(ls == 1.40 for _, ls, _sa in sig_narr) and sig_narr)
    check("narrative: body space-after loosened (10pt)",
          all(sa == 10 for _, _ls, sa in sig_narr))
    check("narrative: text content unchanged",
          [t for t, _, _ in sig_narr] == [t for t, _, _ in sig_none])
    p_after = _render(str(outdir / "after.docx"), None)
    check("no bleed: next render returns to default spacing",
          _spacing_signature(p_after) == sig_none)

    # ---- visual_bias: prose_first flips tiles/body within the section ----
    order_default = _block_order(p_none)
    ws = _ws()
    _write_profile(ws, {"visual_bias": "prose_first"})
    p_prose = _render(str(outdir / "prose_first.docx"), ws)
    order_prose = _block_order(p_prose)

    def _first_body_vs_tbl(order):
        for i, item in enumerate(order):
            if item == "tbl":
                return "tiles"
            if item.startswith("p:First body paragraph"):
                return "body"
        return None

    check("default order: tiles render above the body", _first_body_vs_tbl(order_default) == "tiles")
    check("prose_first: body renders above the tiles", _first_body_vs_tbl(order_prose) == "body")
    check("prose_first: same blocks, order-only change",
          sorted(order_default) == sorted(order_prose))

    # ---- page_cap: warn-only, save proceeds ----
    ws = _ws()
    _write_profile(ws, {"page_cap": {"memo": 1}})
    long_sections = [{"heading": f"Section {i}", "body": ("word " * 400).strip()}
                     for i in range(4)]
    buf = io.StringIO()
    with redirect_stderr(buf):
        bw.make_brief(str(outdir / "overcap.docx"), brief_kind="memo",
                      title="Cap Test", subtitle="sub",
                      exec_header={"verdict": "Cap check."},
                      sections=long_sections, contract="off", voice_gate="off",
                      workspace_root=str(ws))
    check("page_cap: over-cap render still saves (warn-only forever)",
          (outdir / "overcap.docx").exists())
    check("page_cap: stderr carries the [output-profile] note",
          "[output-profile]" in buf.getvalue())
    buf2 = io.StringIO()
    with redirect_stderr(buf2):
        _render(str(outdir / "undercap.docx"), ws)
    check("page_cap: under-cap render is silent",
          "[output-profile]" not in buf2.getvalue())

    # ---- schema registration + writer round-trip ----
    schema = json.loads((ROOT / "shared" / "data-schemas" / "skill_config.schema.json")
                        .read_text(encoding="utf-8"))["skills"]
    check("schema: output_profile registered with the 4 knobs",
          sorted(schema.get("output_profile", [])) ==
          ["default_format", "density", "page_cap", "visual_bias"])
    for s in ("board-pack-assembler", "decision-memo-composer",
              "stress-test", "automation-scanner"):
        check(f"schema: {s} registered (OUT2 §5)", s in schema)
    ws = _ws()
    scw.save_skill_config(ws, "output_profile", {"density": "narrative"})
    check("writer: save_skill_config round-trips into get_output_profile",
          get_output_profile(ws)["density"] == "narrative")
    raised = False
    try:
        scw.save_skill_config(ws, "output_profile", {"densty": "narrative"})  # typo
    except ValueError:
        raised = True
    check("writer: unknown profile key rejected loudly", raised)

    # ---- ⛔ the OUT2 §5 fence: no onboarding surface, no first-run block ----
    onboarding_dir = ROOT / "skills" / "command-room-onboarding"
    hits = []
    for p in onboarding_dir.rglob("*"):
        if p.is_file() and p.suffix in (".md", ".html", ".json", ".yaml", ".py"):
            t = p.read_text(encoding="utf-8", errors="replace").lower()
            if "output_profile" in t or "output profile" in t:
                hits.append(p.name)
    check("fence: onboarding never mentions the output profile", not hits)
    frp = (ROOT / "shared" / "FIRST_RUN_PROTOCOL.md").read_text(encoding="utf-8")
    check("fence: FIRST_RUN_PROTOCOL explicitly excludes it from the catalog",
          "output profile is NOT in this catalog" in frp)

    print()
    if _failures:
        print(f"FAIL — {len(_failures)} output-profile check(s) failed:")
        for f in _failures:
            print(f"  - {f}")
        return 1
    print("ALL SPEC OUT2 §5 output-profile checks PASSED")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
