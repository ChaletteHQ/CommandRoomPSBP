#!/usr/bin/env python3
"""
SPEC OUT2 §2 — shared component library battery.

Covers:
  [1] tile contract — drop-empty refusal, band cap, clean pass (logic moved
      VERBATIM from brief_writer._add_stat_tiles; both backends share it).
  [2] timeline contract — one-point-strip refusal, missing date/label.
  [3] table validation — non-empty, rows-are-lists, column_widths match.
  [4] matrix — dict→grid normalization, star-highlight rule, flag words.
  [5] HTML fragments — tile band + two-column table markup, escaping,
      brand-resolved standalone style (SPEC OUT2 §2c).
  [6] COMPONENT PARITY — the same tile/table input rendered through the
      .docx backend (brief_writer.make_brief) and the HTML backend
      (components.build_*_html) produces the same values/labels/order
      (structural assert, not pixel).
  [7] both backends refuse the same bad input (one contract, two surfaces).

Run via: python3 tests/run_components_test.py
"""
from __future__ import annotations

import html as html_mod
import re
import sys
import tempfile
from pathlib import Path

HERE = Path(__file__).resolve().parent
SCRIPTS = HERE.parent / "shared" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from components import (  # noqa: E402
    MAX_TILES_PER_BAND,
    build_tile_band_html,
    build_two_col_table_html,
    flag_key_for,
    normalize_matrix,
    star_cell_text,
    validate_table,
    validate_tiles,
    validate_timeline,
)

PASS = 0
FAIL = 0


def check(label: str, cond: bool, detail: str = "") -> None:
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  OK {label}")
    else:
        FAIL += 1
        print(f"  FAIL {label}{(' — ' + detail) if detail else ''}")


def raises(fn, *args, **kwargs) -> bool:
    try:
        fn(*args, **kwargs)
        return False
    except ValueError:
        return True


def main() -> int:
    # ------------------------------------------------------------------
    print("[1] tile contract (drop-empty rule, band cap)")
    # ------------------------------------------------------------------
    good = [{"label": "You owe", "value": "2"}, {"label": "Since last touch", "value": "12d"}]
    check("clean band passes", validate_tiles(good) is None)
    check("zero-as-string is data, passes",
          validate_tiles([{"label": "You owe", "value": "0"}]) is None)
    check("empty list refused", raises(validate_tiles, []))
    check("non-list refused", raises(validate_tiles, "tiles"))
    check("empty-value tile refused (drop, never render an empty frame)",
          raises(validate_tiles, [{"label": "You owe", "value": "  "}]))
    check("empty-label tile refused",
          raises(validate_tiles, [{"label": "", "value": "3"}]))
    check("non-dict tile refused", raises(validate_tiles, ["You owe: 2"]))
    over_cap = [{"label": f"L{i}", "value": str(i)} for i in range(MAX_TILES_PER_BAND + 1)]
    check(f"band cap at {MAX_TILES_PER_BAND} enforced", raises(validate_tiles, over_cap))

    # ------------------------------------------------------------------
    print("[2] timeline contract (one-point-strip refusal)")
    # ------------------------------------------------------------------
    tl = [{"date": "Jun 3", "label": "Kickoff call"},
          {"date": "Jul 1", "label": "Deck review", "current": True}]
    check("two-point strip passes", validate_timeline(tl) is None)
    check("one-point strip refused (empty frame)", raises(validate_timeline, tl[:1]))
    check("point without date refused",
          raises(validate_timeline, [{"date": "", "label": "a"}, {"date": "x", "label": "b"}]))
    check("point without label refused",
          raises(validate_timeline, [{"date": "Jun 3", "label": ""}, {"date": "x", "label": "b"}]))

    # ------------------------------------------------------------------
    print("[3] table validation")
    # ------------------------------------------------------------------
    check("n_cols from widest row", validate_table([["a"], ["b", "c"]]) == 2)
    check("n_cols honors wider headers",
          validate_table([["a", "b"]], headers=["1", "2", "3"]) == 3)
    check("empty rows refused", raises(validate_table, []))
    check("non-list row refused", raises(validate_table, ["a b"]))
    check("column_widths mismatch refused",
          raises(validate_table, [["a", "b"]], None, [6.0]))
    check("column_widths match passes",
          validate_table([["a", "b"]], None, [3.0, 3.0]) == 2)

    # ------------------------------------------------------------------
    print("[4] matrix normalization + star + flag words")
    # ------------------------------------------------------------------
    rows, n = normalize_matrix({(0, 0): "x", (1, 2): "y"})
    check("dict cells → dense grid", rows == [["x", "", ""], ["", "", "y"]] and n == 3)
    rows2, n2 = normalize_matrix([["a", "b"]], headers_row=["1", "2", "3"])
    check("list cells pass through; headers widen n_cols", rows2 == [["a", "b"]] and n2 == 3)
    check("empty matrix refused", raises(normalize_matrix, []))
    check("empty dict matrix refused", raises(normalize_matrix, {}))
    check("star applies on star column with content",
          star_cell_text("Option B", 2, 2) == "★ Option B")
    check("star skips other columns", star_cell_text("Option B", 1, 2) == "Option B")
    check("star skips empty cells", star_cell_text("  ", 2, 2) == "  ")
    check("star no-op when star_col_idx None", star_cell_text("x", 0, None) == "x")
    check("flag word → ok key", flag_key_for("Standard") == "flag_ok")
    check("flag synonym with glyphs normalizes", flag_key_for("⚠ Push back!") == "flag_bad")
    check("flag warn synonym", flag_key_for("worth a review") is None
          and flag_key_for("review") == "flag_warn")
    check("non-flag word → None", flag_key_for("Sixty days") is None)
    check("empty → None", flag_key_for("") is None)

    # ------------------------------------------------------------------
    print("[5] HTML fragments (tile band + two-column table)")
    # ------------------------------------------------------------------
    band = build_tile_band_html(good)
    check("band markup uses the widget counter classes",
          band.startswith('<div class="cr-counter-grid">') and band.endswith("</div>"))
    check("band renders labels + values in order",
          band.find("You owe") < band.find("2") < band.find("Since last touch"))
    esc = build_tile_band_html([{"label": "<b>You</b>", "value": 'A "quote" & more'}])
    check("band escapes HTML in labels/values",
          "&lt;b&gt;You&lt;/b&gt;" in esc and "&amp;" in esc and "<b>You</b>" not in esc)
    check("validate=False renders R4-verbatim counters (0 / 6+ buckets ok)",
          build_tile_band_html(
              [{"label": f"B{i}", "value": 0} for i in range(6)], validate=False
          ).count("cr-counter-card") == 6)
    check("validate=True refuses the empty tile",
          raises(build_tile_band_html, [{"label": "x", "value": ""}]))
    check("no <style> by default (widget CSS owns the classes)",
          "<style>" not in band)
    custom = {"palette": {"accent": "8A5A2B"}}
    styled = build_tile_band_html(good, include_style=True,
                                  brand=__import__("brand").get_brand({"workspace": {"brand": custom}}))
    check("standalone style resolves through brand.get_brand (accent override lands)",
          "<style>" in styled and "#8A5A2B" in styled)
    default_styled = build_tile_band_html(good, include_style=True)
    check("standalone default style uses DEFAULT_BRAND values",
          "#2E7D6B" in default_styled and "#102A40" in default_styled)

    table_rows = [["Payment terms", "Net 30"], ["Auto-renewal", "12 months, 60-day notice"]]
    tbl = build_two_col_table_html(table_rows, headers=["Term", "Where it stands"])
    check("two-col table markup", tbl.startswith('<table class="cr-kv-table">')
          and tbl.count("cr-kv-row") == 2 and "cr-kv-th" in tbl)
    check("two-col table renders in order",
          tbl.find("Term") < tbl.find("Payment terms") < tbl.find("Net 30")
          < tbl.find("Auto-renewal"))
    check("two-col: empty rows refused", raises(build_two_col_table_html, []))
    check("two-col: 3-item row refused",
          raises(build_two_col_table_html, [["a", "b", "c"]]))
    check("two-col: bad headers refused",
          raises(build_two_col_table_html, table_rows, headers=["only-one"]))
    check("two-col escapes HTML",
          "&lt;i&gt;" in build_two_col_table_html([["<i>x</i>", "y"]]))

    # ------------------------------------------------------------------
    print("[6] COMPONENT PARITY — same input, docx + HTML backends agree")
    # ------------------------------------------------------------------
    from brief_writer import make_brief  # noqa: E402  (self-installs python-docx)
    from docx import Document  # noqa: E402

    tiles_in = [
        {"label": "Open", "value": "14"},
        {"label": "You owe", "value": "5"},
        {"label": "Oldest", "value": "47d"},
    ]
    out = str(Path(tempfile.mkdtemp(prefix="cr_components_parity_")) / "parity.docx")
    make_brief(
        out,
        brief_kind="past_meeting",
        title="Parity fixture — Sam Sample sync",
        subtitle="Thu, Jul 9 · internal fixture",
        sections=[
            {"heading": "At a glance", "tiles": tiles_in},
            {"heading": "Terms", "table": {
                "rows": table_rows, "headers": ["Term", "Where it stands"]}},
        ],
        contract="off",
        voice_gate="off",
    )
    doc = Document(out)
    # docx tile band = first table (1 row, N cells; cell = [value_p, label_p]).
    band_tbl = doc.tables[0]
    docx_pairs = [
        (cell.paragraphs[1].text, cell.paragraphs[0].text)
        for cell in band_tbl.rows[0].cells
    ]
    html_band = build_tile_band_html(tiles_in)
    html_pairs = [
        (html_mod.unescape(lbl), html_mod.unescape(val))
        for lbl, val in re.findall(
            r'<div class="cr-counter-label">(.*?)</div>'
            r'<div class="cr-counter-value">(.*?)</div>', html_band)
    ]
    check("tile band: same card count", len(docx_pairs) == len(html_pairs) == len(tiles_in))
    check("tile band: same labels, same order (docx small-caps transform aside)",
          [l.lower() for l, _ in docx_pairs] == [l.lower() for l, _ in html_pairs],
          f"docx={docx_pairs} html={html_pairs}")
    check("tile band: same values, same order",
          [v for _, v in docx_pairs] == [v for _, v in html_pairs],
          f"docx={docx_pairs} html={html_pairs}")

    # docx data table = second table (headers row + data rows).
    data_tbl = doc.tables[1]
    docx_table = [[c.text for c in r.cells] for r in data_tbl.rows]
    html_tbl = build_two_col_table_html(table_rows, headers=["Term", "Where it stands"])
    html_headers = [html_mod.unescape(h) for h in re.findall(r'<th class="cr-kv-th">(.*?)</th>', html_tbl)]
    html_rows = [
        [html_mod.unescape(a), html_mod.unescape(b)]
        for a, b in re.findall(
            r'<td class="cr-kv-label">(.*?)</td><td class="cr-kv-value">(.*?)</td>', html_tbl)
    ]
    check("table: same headers, same order", docx_table[0] == html_headers,
          f"docx={docx_table[0]} html={html_headers}")
    check("table: same rows, same values, same order", docx_table[1:] == html_rows,
          f"docx={docx_table[1:]} html={html_rows}")

    # ------------------------------------------------------------------
    print("[7] one contract, two surfaces — both backends refuse the same input")
    # ------------------------------------------------------------------
    bad_tiles = [{"label": "You owe", "value": " "}]
    check("HTML backend refuses the empty tile",
          raises(build_tile_band_html, bad_tiles))
    docx_refused = False
    try:
        make_brief(
            str(Path(tempfile.mkdtemp(prefix="cr_components_refuse_")) / "x.docx"),
            brief_kind="past_meeting", title="t", subtitle="s",
            sections=[{"heading": "h", "tiles": bad_tiles}],
            contract="off", voice_gate="off",
        )
    except ValueError:
        docx_refused = True
    check("docx backend refuses the same empty tile", docx_refused)

    print(f"\n=== Summary: {PASS} passed, {FAIL} failed ===")
    if FAIL:
        return 1
    print("OK — component library battery ALL PASS")
    return 0


if __name__ == "__main__":
    sys.exit(main())
