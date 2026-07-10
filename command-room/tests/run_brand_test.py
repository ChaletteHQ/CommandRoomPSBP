#!/usr/bin/env python3
"""SPEC OUT1 — brand layer + default-theme upgrade + logo plumbing.

Covers (house convention: check(name, cond) prints OK/FAIL, exit 1 on any
failure, auto-discovered by run_all.py; stdlib + jsonschema + python-docx):

  - get_brand() with no config == DEFAULT_BRAND, byte-stable + a fresh copy
  - workspace.brand deep-merge (override wins, unset keys keep the default)
  - orgs[org_id].brand overrides the workspace brand (per-org precedence)
  - validate_brand: default/partial clean; bad hex + unknown key flagged
  - missing logo file -> silent no-logo header (R26); present logo -> letterhead
  - GOLDEN: a no-config brief's theme signature == DEFAULT_BRAND (the deliberate
    default-theme diffs, enumerated) AND the default path (brand=None) matches an
    explicit brand=get_brand() render structurally (defaults are defaults)
  - a branded render restores the default theme globals afterward (no bleed)
  - the entities.schema.json brand $def accepts a valid brand + rejects a bad one
"""
from __future__ import annotations

import base64
import copy
import json
import os
import sys
import tempfile
from pathlib import Path

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
sys.path.insert(0, str(ROOT / "shared" / "scripts"))

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

import brand as brandmod  # noqa: E402
import brief_writer as bw  # noqa: E402
from brand import DEFAULT_BRAND, get_brand, validate_brand  # noqa: E402
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

_failures: list = []


def check(name: str, cond: bool, extra: str = "") -> None:
    print(f"  {'OK  ' if cond else 'FAIL'} {name}{(' — ' + extra) if extra and not cond else ''}")
    if not cond:
        _failures.append(name)


# A minimal valid 1x1 PNG for the letterhead test.
_PNG_1x1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAAC0lEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


def _cell_fill(cell) -> str:
    """The w:fill hex of a table cell's shading, or '' if unshaded."""
    tcpr = cell._tc.tcPr
    if tcpr is None:
        return ""
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        return ""
    return (shd.get(qn("w:fill")) or "").upper()


def _theme_signature(path: str) -> dict:
    """Extract the theme-relevant signature of a rendered brief: title color +
    font, section-heading color, and the tile-band cell fill (if any). Used to
    compare renders without comparing raw bytes (which carry a timestamp)."""
    doc = Document(path)
    sig = {"texts": [p.text for p in doc.paragraphs]}
    # title = the 22pt bold run
    for p in doc.paragraphs:
        if p.runs and p.runs[0].font.size and int(p.runs[0].font.size.pt) == 22:
            r = p.runs[0]
            sig["title_color"] = str(r.font.color.rgb)
            sig["title_font"] = r.font.name
        if p.runs and p.runs[0].font.size and int(p.runs[0].font.size.pt) == 13:
            sig.setdefault("heading_color", str(p.runs[0].font.color.rgb))
    tile_fills = []
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                f = _cell_fill(c)
                if f:
                    tile_fills.append(f)
    sig["fills"] = tile_fills
    return sig


def _base_sections():
    return [{"heading": "Summary", "tiles": [{"label": "Handled", "value": "12"}]},
            {"heading": "Detail", "body": "Some body text with enough words to render fine."}]


def main() -> int:
    print("=== SPEC OUT1: brand layer ===")

    # ---- get_brand() defaults ----
    d = get_brand()
    check("get_brand() == DEFAULT_BRAND", d == DEFAULT_BRAND)
    d["palette"]["accent"] = "000000"
    check("get_brand() returns a fresh copy (no shared mutation)",
          DEFAULT_BRAND["palette"]["accent"] != "000000")

    # ---- workspace.brand deep-merge ----
    ents = {"workspace": {"brand": {"palette": {"accent": "8A5A2B"}, "footer_line": "Acme"}}}
    m = get_brand(ents)
    check("workspace.brand override wins (accent)", m["palette"]["accent"] == "8A5A2B")
    check("workspace.brand keeps unset defaults (ink)",
          m["palette"]["ink"] == DEFAULT_BRAND["palette"]["ink"])
    check("workspace.brand footer merged", m["footer_line"] == "Acme")
    check("workspace.brand keeps default fonts",
          m["fonts"]["heading"] == DEFAULT_BRAND["fonts"]["heading"])

    # ---- org override precedence ----
    ents_org = {
        "workspace": {"brand": {"palette": {"accent": "111111"}}},
        "entities": {"orgs": [{"id": "org_a", "brand": {"palette": {"accent": "222222"}}}]},
    }
    check("org brand overrides workspace brand",
          get_brand(ents_org, "org_a")["palette"]["accent"] == "222222")
    check("workspace-only resolution ignores unrelated org",
          get_brand(ents_org)["palette"]["accent"] == "111111")
    check("unknown org_id falls back to workspace brand",
          get_brand(ents_org, "org_nope")["palette"]["accent"] == "111111")

    # ---- unknown-key safety (a typo can't silently theme) ----
    typo = get_brand({"workspace": {"brand": {"palette": {"acccent": "FF0000"}}}})
    check("palette typo ignored, default accent kept",
          typo["palette"]["accent"] == DEFAULT_BRAND["palette"]["accent"])

    # ---- validate_brand ----
    check("validate: full default clean", validate_brand(DEFAULT_BRAND) == [])
    check("validate: partial clean", validate_brand({"palette": {"accent": "8A5A2B"}}) == [])
    check("validate: bad hex flagged",
          any("accent" in p for p in validate_brand({"palette": {"accent": "#zz"}})))
    check("validate: unknown top key flagged",
          any("foo" in p for p in validate_brand({"foo": 1})))
    check("validate: null logo_path allowed", validate_brand({"logo_path": None}) == [])
    check("validate: non-dict rejected", validate_brand(["nope"]) != [])

    # ---- logo resolution (R26: missing = silent no-logo) ----
    ws = Path(tempfile.mkdtemp(prefix="brand_ws_"))
    (ws / "_hq" / "data").mkdir(parents=True)
    check("missing logo -> None (silent)",
          bw._resolve_logo("_hq/brand/nope.png", str(ws)) is None)
    logo_rel = "_hq/brand/mark.png"
    (ws / "_hq" / "brand").mkdir(parents=True)
    (ws / logo_rel).write_bytes(_PNG_1x1)
    check("present logo -> absolute path",
          bw._resolve_logo(logo_rel, str(ws)) == str(ws / logo_rel))

    outdir = Path(tempfile.mkdtemp(prefix="brand_out_"))

    # default (no-config) render
    p_default = str(outdir / "default.docx")
    bw.make_brief(p_default, brief_kind="memo", title="Theme Test", subtitle="sub",
                  exec_header={"verdict": "Theme check."},  # OUT2 §4 flip: memo requires it
                  sections=_base_sections(), contract="off", voice_gate="off")
    sig_default = _theme_signature(p_default)

    # GOLDEN: deliberate default-theme diffs are exactly DEFAULT_BRAND's values.
    check("golden: title color == default heading navy",
          sig_default["title_color"] == DEFAULT_BRAND["palette"]["heading"])
    check("golden: title font == default heading font (Georgia)",
          sig_default["title_font"] == DEFAULT_BRAND["fonts"]["heading"])
    check("golden: heading color == default heading navy",
          sig_default.get("heading_color") == DEFAULT_BRAND["palette"]["heading"])
    check("golden: tile band uses default tile_bg fill",
          DEFAULT_BRAND["palette"]["tile_bg"].upper() in sig_default["fills"])

    # explicit brand=get_brand() must match the None-path render structurally.
    p_explicit = str(outdir / "explicit.docx")
    bw.make_brief(p_explicit, brief_kind="memo", title="Theme Test", subtitle="sub",
                  exec_header={"verdict": "Theme check."},  # OUT2 §4 flip: memo requires it
                  sections=_base_sections(), contract="off", voice_gate="off",
                  brand=get_brand())
    check("defaults: brand=None render == explicit brand=get_brand() render",
          _theme_signature(p_explicit) == sig_default)

    # ---- branded render + default restoration (no bleed) ----
    p_branded = str(outdir / "branded.docx")
    bw.make_brief(p_branded, brief_kind="memo", title="Theme Test", subtitle="sub",
                  exec_header={"verdict": "Theme check."},  # OUT2 §4 flip: memo requires it
                  sections=_base_sections(), contract="off", voice_gate="off",
                  brand=get_brand({"workspace": {"brand": {"palette": {"heading": "AA0000"}}}}))
    sig_branded = _theme_signature(p_branded)
    check("branded: title color picked up the brand override",
          sig_branded["title_color"] == "AA0000")
    check("no bleed: default theme restored after a branded render",
          bw.HEADING_FONT == DEFAULT_BRAND["fonts"]["heading"]
          and str(bw.HEADING) == DEFAULT_BRAND["palette"]["heading"])

    # a render with the workspace brand resolves via workspace_root; delete it and
    # the next render returns to defaults (acceptance #3).
    entities_path = ws / "_hq" / "data" / "entities.json"
    entities_path.write_text(json.dumps({
        "version": 1, "entities": {"people": [], "projects": [], "orgs": []},
        "workspace": {"brand": {"palette": {"heading": "BB0000"}, "logo_path": logo_rel}},
    }), encoding="utf-8")
    p_ws = str(outdir / "ws_brand.docx")
    bw.make_brief(p_ws, brief_kind="memo", title="Theme Test", subtitle="sub",
                  exec_header={"verdict": "Theme check."},  # OUT2 §4 flip: memo requires it
                  sections=_base_sections(), contract="off", voice_gate="off",
                  workspace_root=str(ws))
    sig_ws = _theme_signature(p_ws)
    check("acceptance #2: workspace.brand picked up via workspace_root",
          sig_ws["title_color"] == "BB0000")
    check("acceptance #2: workspace logo rendered a letterhead image",
          any("graphic" in p._p.xml.lower() or "drawing" in p._p.xml.lower()
              for p in Document(p_ws).paragraphs[:2]))
    # delete the brand object
    entities_path.write_text(json.dumps({
        "version": 1, "entities": {"people": [], "projects": [], "orgs": []},
        "workspace": {},
    }), encoding="utf-8")
    p_ws2 = str(outdir / "ws_nobrand.docx")
    bw.make_brief(p_ws2, brief_kind="memo", title="Theme Test", subtitle="sub",
                  exec_header={"verdict": "Theme check."},  # OUT2 §4 flip: memo requires it
                  sections=_base_sections(), contract="off", voice_gate="off",
                  workspace_root=str(ws))
    check("acceptance #3: deleting the brand returns output to defaults",
          _theme_signature(p_ws2) == sig_default)

    # ---- schema: brand $def accepts valid, rejects bad ----
    try:
        import jsonschema
        schema = json.loads((ROOT / "shared" / "data-schemas" / "entities.schema.json")
                            .read_text(encoding="utf-8"))
        brand_schema = {"$defs": schema["$defs"], "$ref": "#/$defs/brand"}
        V = jsonschema.Draft202012Validator(brand_schema)
        check("schema: valid brand passes",
              list(V.iter_errors({"palette": {"accent": "8A5A2B"}, "fonts": {"heading": "Georgia"}})) == [])
        check("schema: bad hex rejected",
              len(list(V.iter_errors({"palette": {"accent": "#zz"}}))) > 0)
        check("schema: unknown palette key rejected",
              len(list(V.iter_errors({"palette": {"nope": "FFFFFF"}}))) > 0)
    except ImportError:
        print("  SKIP jsonschema unavailable — schema checks skipped")

    print()
    if _failures:
        print(f"FAIL — {len(_failures)} brand check(s) failed:")
        for f in _failures:
            print(f"  - {f}")
        return 1
    print("ALL SPEC OUT1 brand checks PASSED")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
