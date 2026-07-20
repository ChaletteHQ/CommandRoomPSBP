#!/usr/bin/env python3
"""
Polished brief docx writer (v2.14.32+).

Single source of truth for Command Room meeting brief layout. Both Call_Prep
and Past_Meeting briefs go through `make_brief()`. Replaces the v2.14.31-era
flow where each fire asked the docx skill to lay out the document and got
slightly different output every time.

Save-time gates (canonical order): input validation → output-contract gate
(SPEC B3, PRE-Document()) → voice-tell gate (SPEC B2, PRE-Document(), see
`make_brief`) → render → post-render leak scan (`docx_leak_scanner`). The
contract and voice gates both raise before any file is written; the leak scan
runs after save.

Design goals:
  - Deterministic. Same structured input -> same docx output, every fire.
  - Forwardable-clean per `meeting-notes/SKILL.md` Brief Authoring Rules.
    No provenance footers, no internal IDs, no calendar URLs.
  - Visually polished but quiet. Subtle colors, generous whitespace,
    a single horizontal rule under the header. No emoji, no logos,
    no loud branding.
  - One US Letter page friendly (margins + font sizes tuned for density).

Used by:
  - orchestrator-upcoming-meetings.md (Phase 4 step 3 -> Call_Prep_*.docx)
  - orchestrator-past-meetings.md (Phase 4 step 7 -> Past_Meeting_*.docx)
  - meeting-notes/SKILL.md Step 9a (on-demand Past_Meeting_*.docx)

Dependencies: python-docx (pinned to PYTHON_DOCX_PIN below). Self-installs
on import failure (idempotent, ~3s cold, no-op when already present). The
pin must match what the brief_writer test suite is run against — see
README "Requirements" section for how to pre-install in locked-down
environments.
"""
from __future__ import annotations

import json
import re
import subprocess
import sys
from pathlib import Path
from typing import List, Dict, Optional, Union

_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))


PYTHON_DOCX_PIN = "1.2.0"


def _ensure_python_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        print(
            f"Installing python-docx (=={PYTHON_DOCX_PIN}) — one-time setup. "
            "(Plugin requires this for .docx generation. "
            "See README Requirements section to pre-install in locked-down environments.)",
            file=sys.stderr,
        )
        subprocess.run(
            [
                sys.executable,
                "-m",
                "pip",
                "install",
                "--quiet",
                f"python-docx=={PYTHON_DOCX_PIN}",
            ],
            check=True,
        )


_ensure_python_docx()

from docx import Document  # noqa: E402
from docx.shared import Pt, Inches, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402


# ---------- Brand theme (SPEC OUT1 — resolved via shared/scripts/brand.py) ----------
# Every color / font / footer / logo constant is now sourced from the brand
# layer instead of being hardcoded here. `get_brand()` with no config returns
# DEFAULT_BRAND (an UPGRADED quiet-professional default — the deliverable), so a
# fresh workspace looks premium with zero config. A paying client's `brand`
# object in entities.json (workspace-level or per-org) deep-merges over the
# default; an absent brand object = byte-stable defaults, no warning, no event.
#
# The module-level globals below hold the CURRENT render theme. They default to
# DEFAULT_BRAND at import and are re-applied per render by `_apply_brand()` (the
# `brand=` / workspace_root / org_id resolution in make_brief), then restored.
# All the _add_* helpers read these globals, so a per-render theme flows through
# without threading a palette arg into every helper.

from brand import get_brand  # noqa: E402

# SPEC OUT2 §2a — the component data shapes + validation (tiles / timeline /
# table / matrix drop-empty rules, one-point-strip refusal, star-highlight,
# flag words) moved to the shared component library. This file is now the
# .docx BACKEND consuming those shapes; the HTML fragment backend lives in
# components.py too (build_tile_band_html / build_two_col_table_html).
from components import (  # noqa: E402
    validate_tiles as _validate_tiles,
    validate_timeline as _validate_timeline,
    validate_table as _validate_table,
    normalize_matrix as _normalize_matrix,
    star_cell_text as _star_cell_text,
    flag_key_for as _flag_key_for,
)

# SPEC OUT2 §5 — the cross-skill output profile (density / visual_bias /
# page_cap / default_format). Resolved per render exactly like the brand:
# absent file = DEFAULT_OUTPUT_PROFILE = byte-identical to pre-profile output.
from output_profile import (  # noqa: E402
    DEFAULT_OUTPUT_PROFILE,
    get_output_profile,
)

# ---------- The shared gate stack (SPEC OUT5 §3b) ----------
# The kind registry, EXEC1 kind sets, gate sequence, and audit emitters moved
# VERBATIM to brief_gates.py so the premium-HTML backend (premium_html.py) runs
# the IDENTICAL stack — the parity invariant is pinned by
# tests/run_guard_g16_gate_parity_test.py. Re-exported here so every existing
# import (`from brief_writer import STANDARD_KINDS`, the output-contract
# validator's sync guard on EYEBROW_BY_KIND, …) keeps working unchanged.
from brief_gates import (  # noqa: E402
    EYEBROW_BY_KIND,
    SUPPORTED_BRIEF_KINDS,
    STANDARD_KINDS,
    DECISION_SHAPED_KINDS,
    EXEC_EYEBROW_EXCLUDED_KINDS,
    ASKS_HEADING,
    MAX_ASKS,
    EXEC_HEADER_LINES as _EXEC_HEADER_LINES,
    run_pre_save_gates as _run_pre_save_gates,
    warn_page_cap as _warn_page_cap,
    estimate_pages as _estimate_pages_shared,
    emit_brief_meta_audit as _emit_brief_meta_audit,
    emit_gate_ran_audit as _emit_gate_ran_audit,
)

_DEFAULT_RESOLVED = get_brand()  # pure defaults; no I/O at import


def _rgb(hexstr: str) -> "RGBColor":
    """A 6-hex string (no '#') -> RGBColor."""
    h = str(hexstr).lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# Current-render theme globals (rebound by _apply_brand; the _add_* helpers read
# these). Colors that go through _set_run are RGBColor; colors that go through
# _shade_cell stay bare hex strings (its API wants hex).
INK = _rgb(_DEFAULT_RESOLVED["palette"]["ink"])
HEADING = _rgb(_DEFAULT_RESOLVED["palette"]["heading"])
MUTED = _rgb(_DEFAULT_RESOLVED["palette"]["muted"])
ACCENT = _rgb(_DEFAULT_RESOLVED["palette"]["accent"])
RULE_HEX = _DEFAULT_RESOLVED["palette"]["rule"]

TILE_BG = _DEFAULT_RESOLVED["palette"]["tile_bg"]
ZEBRA = _DEFAULT_RESOLVED["palette"]["zebra"]
TABLE_HEADER = _DEFAULT_RESOLVED["palette"]["table_header"]
COL_HEADER = _DEFAULT_RESOLVED["palette"]["col_header"]
HIGHLIGHT = _DEFAULT_RESOLVED["palette"]["highlight"]
FLAG_OK = _DEFAULT_RESOLVED["palette"]["flag_ok"]
FLAG_WARN = _DEFAULT_RESOLVED["palette"]["flag_warn"]
FLAG_BAD = _DEFAULT_RESOLVED["palette"]["flag_bad"]

BODY_FONT = _DEFAULT_RESOLVED["fonts"]["body"]
HEADING_FONT = _DEFAULT_RESOLVED["fonts"]["heading"]
MONO_FONT = _DEFAULT_RESOLVED["fonts"]["mono"]

FOOTER_DEFAULT = _DEFAULT_RESOLVED["footer_line"]
LOGO_PATH = _DEFAULT_RESOLVED["logo_path"]
EYEBROW_STYLE = dict(_DEFAULT_RESOLVED["eyebrow_style"])

# Current-render output-profile globals (SPEC OUT2 §5 — rebound per render by
# _apply_output_profile, restored in make_brief's finally, exactly like the
# brand globals). Defaults equal today's behavior byte-for-byte:
#   tight     -> body line_spacing 1.25 / space_after 6pt (the pre-profile values)
#   narrative -> body line_spacing 1.40 / space_after 10pt (looser prose)
_BODY_LINE_SPACING = 1.25
_BODY_SPACE_AFTER = 6
_VISUAL_BIAS = DEFAULT_OUTPUT_PROFILE["visual_bias"]

def _apply_output_profile(profile: dict) -> None:
    """Rebind the render-density globals from a resolved output profile.
    Called at the top of every render and restored to defaults after
    (make_brief try/finally). Single-threaded per fire, same posture as
    _apply_brand."""
    global _BODY_LINE_SPACING, _BODY_SPACE_AFTER, _VISUAL_BIAS
    if profile.get("density") == "narrative":
        _BODY_LINE_SPACING = 1.40
        _BODY_SPACE_AFTER = 10
    else:  # "tight" — today's behavior
        _BODY_LINE_SPACING = 1.25
        _BODY_SPACE_AFTER = 6
    _VISUAL_BIAS = (
        profile.get("visual_bias")
        if profile.get("visual_bias") in ("tiles_first", "prose_first")
        else "tiles_first"
    )


# SPEC OUT5 — the page estimate moved to brief_gates.estimate_pages (shared so
# a configured cap warns identically on both backends). Alias kept for any
# in-repo caller of the old name.
_estimate_pages = _estimate_pages_shared

# Flag-cell tint words for the contract-review matrix (SPEC OUT1 §4) moved to
# components.FLAG_TINT_KEYS (SPEC OUT2 §2a) — this backend maps the returned
# palette KEY to its resolved brand tint hex in _flag_tint_for below.


def _apply_brand(brand: dict) -> None:
    """Rebind the module-level render-theme globals from a resolved brand dict.
    Called at the top of every render with the render's resolved theme, and
    restored to defaults after (make_brief try/finally). Single-threaded per
    fire, so global rebinding is safe and keeps the _add_* helpers arg-free."""
    global INK, HEADING, MUTED, ACCENT, RULE_HEX
    global TILE_BG, ZEBRA, TABLE_HEADER, COL_HEADER, HIGHLIGHT
    global FLAG_OK, FLAG_WARN, FLAG_BAD
    global BODY_FONT, HEADING_FONT, MONO_FONT
    global FOOTER_DEFAULT, LOGO_PATH, EYEBROW_STYLE
    p = brand["palette"]
    INK = _rgb(p["ink"])
    HEADING = _rgb(p["heading"])
    MUTED = _rgb(p["muted"])
    ACCENT = _rgb(p["accent"])
    RULE_HEX = p["rule"]
    TILE_BG = p["tile_bg"]
    ZEBRA = p["zebra"]
    TABLE_HEADER = p["table_header"]
    COL_HEADER = p["col_header"]
    HIGHLIGHT = p["highlight"]
    FLAG_OK = p["flag_ok"]
    FLAG_WARN = p["flag_warn"]
    FLAG_BAD = p["flag_bad"]
    BODY_FONT = brand["fonts"]["body"]
    HEADING_FONT = brand["fonts"]["heading"]
    MONO_FONT = brand["fonts"]["mono"]
    FOOTER_DEFAULT = brand["footer_line"]
    LOGO_PATH = brand["logo_path"]
    EYEBROW_STYLE = dict(brand["eyebrow_style"])




# ---------- Internal helpers ----------

def _set_run(run, *, font=None, size=11, color=None, bold=False, italic=False):
    # Resolve None -> current render-theme global at CALL time, so a per-render
    # brand applied via _apply_brand reaches callers that omit font/color (def
    # defaults bind once at import and would otherwise miss a rebind).
    run.font.name = font if font is not None else BODY_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color if color is not None else INK
    run.font.bold = bold
    run.font.italic = italic


def _add_horizontal_rule(paragraph) -> None:
    """Single light-grey rule below the given paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_normal_baseline(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK


def _tighten_margins(doc) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _add_eyebrow(doc, label: str) -> None:
    style = EYEBROW_STYLE or {}
    text = label.upper() if style.get("upper", True) else label
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, font=HEADING_FONT, size=style.get("size", 9),
             color=ACCENT, bold=style.get("bold", True))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)


def _add_title(doc, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    _set_run(run, font=HEADING_FONT, size=22, color=HEADING, bold=True)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15


def _add_subtitle(doc, subtitle: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    _set_run(run, font=BODY_FONT, size=11, color=MUTED)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)


def _add_header_rule(doc) -> None:
    p = doc.add_paragraph()
    _add_horizontal_rule(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)


def _add_section_heading(doc, text: str, level: int = 2) -> None:
    """v3.13.8+ — uses python-docx's actual Heading{level} style so Word's
    TOC, accessibility tooling, and outline panes see real headings (Bug #7).

    Pre-v3.13.8 this function rendered headings as bold runs inside Normal
    paragraphs, which looked right but didn't register as headings in
    Word's structure. Now we apply the Heading2 style and then override
    font/color so the canonical navy + Calibri 13pt look stays intact.
    """
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    _set_run(run, font=HEADING_FONT, size=13, color=HEADING, bold=True)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True


def _add_body_paragraphs(doc, body_text: str) -> None:
    """Split on blank lines; each block becomes one paragraph."""
    for block in body_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph()
        run = p.add_run(block)
        _set_run(run, font=BODY_FONT, size=11, color=INK)
        p.paragraph_format.space_before = Pt(0)
        # SPEC OUT2 §5 — density comes from the output profile ("tight" default
        # keeps the pre-profile 6pt/1.25 values byte-stably).
        p.paragraph_format.space_after = Pt(_BODY_SPACE_AFTER)
        p.paragraph_format.line_spacing = _BODY_LINE_SPACING


def _shade_cell(cell, hex_color: str) -> None:
    """Apply a background shading color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_table(
    doc,
    rows: List[List[str]],
    headers: Optional[List[str]] = None,
    highlight_row_idx: Optional[int] = None,
    column_widths: Optional[List[float]] = None,
) -> None:
    """v3.13.8+ — table primitive for brief_writer.

    Args:
      rows: list of lists — table data rows (each inner list = one row).
      headers: optional list — header row content. When provided, gets
        navy heading style + bold white text + slightly heavier height.
      highlight_row_idx: optional int — row index within `rows` (0-based,
        excluding the header) to visually highlight, e.g. the recommended
        option in a comparison matrix. Renders with a subtle ACCENT-tinted
        background.
      column_widths: optional list of float (inches). Defaults to equal split
        across 6 inches usable page width.
    """
    # SPEC OUT2 §2a — shape validation shared with the HTML backend
    # (components.validate_table; messages + rules identical to pre-OUT2).
    n_cols = _validate_table(rows, headers, column_widths)

    total_rows = len(rows) + (1 if headers else 0)
    table = doc.add_table(rows=total_rows, cols=n_cols)
    table.style = "Table Grid"

    # Column widths
    if column_widths:
        for i, width in enumerate(column_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    else:
        per_col = 6.0 / n_cols
        for i in range(n_cols):
            for row in table.rows:
                row.cells[i].width = Inches(per_col)

    # Header row
    row_offset = 0
    if headers:
        for j, header_text in enumerate(headers):
            cell = table.rows[0].cells[j]
            _shade_cell(cell, TABLE_HEADER)  # heading-navy header fill
            for p in cell.paragraphs:
                p.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(header_text))
            _set_run(run, font=HEADING_FONT, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
        row_offset = 1

    # Data rows
    for i, row_data in enumerate(rows):
        target_row = table.rows[i + row_offset]
        is_highlight = highlight_row_idx is not None and i == highlight_row_idx
        # Zebra stripe (very subtle) when no per-row highlight has consumed it
        zebra_fill = ZEBRA if (i % 2 == 1) else None
        accent_fill = HIGHLIGHT  # ACCENT-tinted background
        for j in range(n_cols):
            cell = target_row.cells[j]
            text = str(row_data[j]) if j < len(row_data) else ""
            if is_highlight:
                _shade_cell(cell, accent_fill)
            elif zebra_fill:
                _shade_cell(cell, zebra_fill)
            for p in cell.paragraphs:
                p.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            _set_run(
                run,
                font=BODY_FONT,
                size=10,
                color=INK,
                bold=is_highlight,
            )

    # Space after the table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def _flag_tint_for(value: str) -> Optional[str]:
    """The palette tint hex for a flag cell's value, or None if it doesn't read
    as a flag word. Word normalization lives in components.flag_key_for (SPEC
    OUT2 §2a); this backend maps the palette KEY to the CURRENT render theme's
    tint hex. Used by the contract-review matrix so a green/amber/red flag
    column renders as shaded cells, not prose labels."""
    key = _flag_key_for(value)
    if key is None:
        return None
    return {"flag_ok": FLAG_OK, "flag_warn": FLAG_WARN, "flag_bad": FLAG_BAD}[key]


def _add_matrix(
    doc,
    cells: Union[List[List[str]], Dict[tuple, str]],
    headers_row: Optional[List[str]] = None,
    headers_col: Optional[List[str]] = None,
    star_col_idx: Optional[int] = None,
    flag_col_idx: Optional[int] = None,
) -> None:
    """v3.13.8+ — N×M comparison matrix with optional ★ glyph highlighting.

    Args:
      cells: either a 2D list of strings OR a dict {(row, col): value}.
      headers_row: optional column-header labels (top row).
      headers_col: optional row-header labels (leftmost column).
      star_col_idx: optional int — render ★ glyph in front of every cell
        whose row's value in this column is non-empty/non-falsy (used for
        "recommended option" matrices).
      flag_col_idx: optional int (SPEC OUT1 §4) — a column of flag words
        (ok / warn / bad and common synonyms) whose cells get shaded with the
        brand flag tint that matches each cell's value. Used by contract-review
        so the green/yellow/red flag column is color, not just a prose label.
    """
    # SPEC OUT2 §2a — dict→grid normalization + shape validation shared with
    # the HTML backend (components.normalize_matrix; logic identical).
    rows_data, n_cols = _normalize_matrix(cells, headers_row)
    # +1 column for headers_col if present
    actual_cols = n_cols + (1 if headers_col else 0)
    total_rows = len(rows_data) + (1 if headers_row else 0)
    table = doc.add_table(rows=total_rows, cols=actual_cols)
    table.style = "Table Grid"

    per_col = 6.0 / actual_cols
    for i in range(actual_cols):
        for row in table.rows:
            row.cells[i].width = Inches(per_col)

    # Header row
    row_offset = 0
    if headers_row:
        if headers_col:
            cell = table.rows[0].cells[0]
            _shade_cell(cell, TABLE_HEADER)
        for j, header_text in enumerate(headers_row):
            target_j = j + (1 if headers_col else 0)
            cell = table.rows[0].cells[target_j]
            _shade_cell(cell, TABLE_HEADER)
            for p in cell.paragraphs:
                p.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(header_text))
            _set_run(run, font=HEADING_FONT, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
        row_offset = 1

    # Data + headers_col
    for i, row_data in enumerate(rows_data):
        target_row = table.rows[i + row_offset]
        if headers_col:
            label = headers_col[i] if i < len(headers_col) else ""
            cell = target_row.cells[0]
            _shade_cell(cell, COL_HEADER)
            for p in cell.paragraphs:
                p.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(label))
            _set_run(run, font=HEADING_FONT, size=10, color=HEADING, bold=True)
        for j in range(n_cols):
            target_j = j + (1 if headers_col else 0)
            cell = target_row.cells[target_j]
            text = str(row_data[j]) if j < len(row_data) else ""
            # SPEC OUT2 §2a — star-highlight rule shared with the HTML backend.
            text = _star_cell_text(text, j, star_col_idx)
            if flag_col_idx is not None and j == flag_col_idx:
                tint = _flag_tint_for(text)
                if tint:
                    _shade_cell(cell, tint)
            for p in cell.paragraphs:
                p.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            _set_run(run, font=BODY_FONT, size=10, color=INK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def _add_stat_tiles(doc, tiles: List[Dict[str, str]]) -> None:
    """v4.5.2 S1 — the stat-tile band (visual layer, M directive).

    Renders 1-5 tiles as a single borderless row: big value over a small-caps
    accent label. REFUSES an empty tile — the drop rule ("a tile with no data
    is dropped, never rendered as an empty frame") is enforced by the builder
    (prep_pipeline.build_prep_tiles) AND here at the render chokepoint, so an
    empty frame is structurally impossible whichever path called us.

    SPEC OUT2 §2a — the validation moved to components.validate_tiles (shared
    with the HTML backend); rules + messages identical.
    """
    _validate_tiles(tiles)

    table = doc.add_table(rows=1, cols=len(tiles))
    per_col = 6.0 / len(tiles)
    for i in range(len(tiles)):
        table.rows[0].cells[i].width = Inches(per_col)
    for i, t in enumerate(tiles):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, TILE_BG)
        for p in cell.paragraphs:
            p.text = ""
        value_p = cell.paragraphs[0]
        value_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = value_p.add_run(str(t["value"]).strip())
        _set_run(run, font=HEADING_FONT, size=16, color=HEADING, bold=True)
        value_p.paragraph_format.space_before = Pt(4)
        value_p.paragraph_format.space_after = Pt(0)
        label_p = cell.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = label_p.add_run(str(t["label"]).strip().upper())
        _set_run(run, font=HEADING_FONT, size=8, color=ACCENT, bold=True)
        label_p.paragraph_format.space_before = Pt(0)
        label_p.paragraph_format.space_after = Pt(4)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def _add_timeline(doc, points: List[Dict[str, str]]) -> None:
    """v4.5.2 S1 — the relationship timeline strip (visual layer).

    Compact one-line-per-point strip: accent date + label, the current
    meeting bold with an em-dash marker. Requires >= 2 points — a one-point
    strip is an empty frame; the builder (prep_pipeline) already drops it,
    and this render-chokepoint check makes the rule bypass-proof.

    SPEC OUT2 §2a — the validation moved to components.validate_timeline
    (shared with the HTML backend); rules + messages identical.
    """
    _validate_timeline(points)

    for pt in points:
        is_current = bool(pt.get("current"))
        p = doc.add_paragraph()
        date_run = p.add_run(str(pt["date"]).strip() + "   ")
        _set_run(date_run, font=HEADING_FONT, size=9, color=ACCENT, bold=True)
        label_text = str(pt["label"]).strip()
        if is_current:
            label_text += "  — this meeting"
        label_run = p.add_run(label_text)
        _set_run(label_run, font=BODY_FONT, size=9,
                 color=INK if is_current else MUTED, bold=is_current)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def _add_bullets(doc, bullets: List[str]) -> None:
    for line in bullets:
        line = line.strip()
        if not line:
            continue
        # Strip leading bullet markers if the caller already added them.
        if line.startswith(("- ", "* ", "• ")):
            line = line[2:].lstrip()
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(line)
        _set_run(run, font=BODY_FONT, size=11, color=INK)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.20


def _add_footer(doc, text: str) -> None:
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(text)
    _set_run(run, font=BODY_FONT, size=9, color=MUTED)


def _resolve_logo(logo_path: Optional[str], workspace_root: Optional[str]) -> Optional[str]:
    """Absolute path of the brand logo IF it exists on disk, else None.

    CONTRACT R26 / privacy (SPEC OUT1 §3d): a configured-but-missing logo NEVER
    raises in a client chat — it silently falls back to the quiet no-logo
    header. An absolute path is used as-is; a relative path resolves against
    workspace_root. `None` (the default) means no logo at all."""
    if not logo_path:
        return None
    p = Path(logo_path)
    if not p.is_absolute() and workspace_root:
        p = Path(workspace_root) / logo_path
    try:
        return str(p) if p.is_file() else None
    except OSError:
        return None


def _add_logo_header(doc, logo_abspath: str) -> None:
    """Right-aligned letterhead image at the very top of the document, <= 0.35in
    tall (SPEC OUT1 §3b). Only called when a brand logo file actually exists;
    the default (no logo) leaves the current quiet header untouched. Any
    python-docx image error is swallowed — a bad image never blocks a client's
    deliverable (R26)."""
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(logo_abspath, height=Inches(0.35))
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
    except Exception:
        # Never let a malformed / unreadable image block the save.
        pass


# FS-13 — EXEC_EYEBROW_EXCLUDED_KINDS now lives in brief_gates (imported above)
# so both backends share the verdict-only rule for document/decision kinds.


def _add_exec_header(doc, exec_header: Dict[str, str], brief_kind: str = "") -> None:
    """SPEC EXEC1 element 1 — the 30-second contract. Brief-family kinds render
    the full block:

      [bold verdict sentence]
      CHANGED  ...
      DECIDE   ...
      NEEDED   ...
      <light rule>

    Document / decision kinds (EXEC_EYEBROW_EXCLUDED_KINDS — FS-13) render the
    VERDICT lead + rule ONLY; the CHANGED/DECIDE/NEEDED eyebrow is a brief
    scaffold and does not belong on them.

    `exec_header` keys: verdict (required), changed, decide, needs. Any of the
    three eyebrow lines may be a 'nothing-form' per the anti-washing floor —
    the renderer is content-agnostic; the concreteness floor is a
    checklist/validator concern, not a render concern.
    """
    verdict = (exec_header.get("verdict") or "").strip()
    if verdict:
        p = doc.add_paragraph()
        run = p.add_run(verdict)
        _set_run(run, font=HEADING_FONT, size=12, color=INK, bold=True)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

    if brief_kind not in EXEC_EYEBROW_EXCLUDED_KINDS:
        for key, label in _EXEC_HEADER_LINES:
            text = (exec_header.get(key) or "").strip()
            if not text:
                continue
            p = doc.add_paragraph()
            label_run = p.add_run(label + "  ")
            _set_run(label_run, font=HEADING_FONT, size=9, color=ACCENT, bold=True)
            text_run = p.add_run(text)
            _set_run(text_run, font=BODY_FONT, size=11, color=INK)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.20

    _add_header_rule(doc)


def _add_charts(doc, charts_spec: List[dict], resolved_brand: dict) -> None:
    """SPEC OUT3 — best-effort chart images under the tile band.

    Each entry is a {kind, data, title?} spec for `charts.try_chart_png` (the
    render chokepoint: build_chart -> rasterize ladder). A shape refusal, a
    leak finding, or a missing rasterizer on this machine returns None and the
    chart is simply skipped — the section's table/tile representation of the
    same numbers stands, byte-identical to pre-OUT3 (the visual_gate posture:
    upgrade machines that can render, never degrade ones that can't). A
    python-docx image error is swallowed like _add_logo_header's (R26). The
    charts module is imported lazily so a partial install can never take the
    brief writer down (the FS-15 lesson; same posture as the leak scanner)."""
    try:
        from charts import try_chart_png
    except ImportError:
        return
    for spec in charts_spec:
        png = try_chart_png(spec, brand=resolved_brand)
        if not png:
            continue
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(png, width=Inches(6.0))
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
        except Exception:
            # Never let a malformed image block the save (R26).
            pass


def _add_asks_block(doc, asks: List[Dict[str, str]]) -> None:
    """SPEC EXEC1 element 4 — the ASK block. Rendered as the LAST content block
    (after sections, before the footer), under the canonical heading. Zero asks
    → caller skips this entirely (the exec header already said NEEDED Nothing).
    Each ask is a reader-action; an optional deadline renders as ' — by <when>'."""
    _add_section_heading(doc, ASKS_HEADING)
    lines = []
    for ask in asks:
        text = (ask.get("text") or "").strip()
        if not text:
            continue
        deadline = (ask.get("deadline") or "").strip()
        lines.append(f"{text} — by {deadline}" if deadline else text)
    _add_bullets(doc, lines)


# ---------- Public API ----------

def make_brief(
    output_path: str,
    *,
    brief_kind: str,
    title: str,
    subtitle: str,
    sections: List[Dict[str, Union[str, List[str]]]],
    footer_text: Optional[str] = None,
    voice_gate: str = "default",
    contract: str = "enforce",
    contract_profile: Optional[str] = None,
    exec_header: Optional[Dict[str, str]] = None,
    asks: Optional[List[Dict[str, str]]] = None,
    workspace_root: Optional[str] = None,
    brand: Optional[dict] = None,
    org_id: Optional[str] = None,
) -> str:
    """Write a polished brief .docx to `output_path` and return the path.

    Args:
      output_path: absolute path to write the .docx (already resolved via
        `brief_path.get_brief_path()`).
      brief_kind: one of the supported kinds. Drives the eyebrow label
        only — layout is identical across kinds. Pre-v3.10.0 only
        "call_prep" and "past_meeting" were supported; v3.10.0+ extends
        the set to cover the v3.8.0 new-skill deliverables.

        Supported values:
          "call_prep"          → eyebrow "CALL PREP"
          "past_meeting"       → eyebrow "MEETING BRIEF"
          "board_pack"         → eyebrow "BOARD PACK"
          "contract_review"    → eyebrow "CONTRACT REVIEW"
          "decision_memo"      → eyebrow "DECISION MEMO"
          "operator_report"    → eyebrow "OPERATING LIFT"
          "value_receipt"      → eyebrow "VALUE RECEIPT"
          "weekly_recap"       → eyebrow "WEEKLY RECAP"
          "weekly_audit"       → eyebrow "WEEKLY AUDIT"
          "dormant_scan"       → eyebrow "DORMANT CUSTOMER SCAN"
          "automation_scan"    → eyebrow "AUTOMATION SCAN"
          "automation_recipe"  → eyebrow "AUTOMATION SETUP RECIPE"
          "followup_pack"      → eyebrow "FOLLOW-UP PACK"
          "memo"               → eyebrow "MEMO"
          "one_pager"          → eyebrow "ONE-PAGER"
          "insights"           → eyebrow "INSIGHTS"
          "stress_test"        → eyebrow "STRESS TEST"
          "kpi_scorecard"      → eyebrow "KPI SCORECARD"  (SPEC OUT7)
      title: top-of-page title. Should resolve all entity IDs to names
        (e.g. "Sam Sample — Q2 deck review", not "person_037 — ...").
      subtitle: one-line meta below title (e.g. "Fri, May 9 · 9:00 AM PT
        · Summit Company"). Plain text, no formatting.
      sections: ordered list of section dicts. Each dict has:
          - "heading": str — section title
          - "body": str (optional) — paragraph-style content; blank lines
            become paragraph breaks
          - "bullets": list[str] (optional) — bullet items
          - "table": dict (optional) — {rows: list[list], headers?: list,
            highlight_row_idx?: int, column_widths?: list[float]}
          - "matrix": dict (optional) — {cells: 2D list or {(r,c): v} dict,
            headers_row?: list, headers_col?: list, star_col_idx?: int}
          - "tiles": list[{label, value}] (optional, v4.5.2 S1) — stat-tile
            band (1-5 tiles). Empty tiles are REFUSED — the caller drops a
            tile with no data (prep_pipeline.build_prep_tiles), never renders
            an empty frame.
          - "timeline": list[{date, label, current?}] (optional, v4.5.2 S1) —
            relationship timeline strip; >= 2 points required (the caller
            drops the section below that).
          - "charts": list[{kind, data, title?}] (optional, SPEC OUT3) —
            substrate-derived charts via shared/scripts/charts.py. Rendered
            best-effort at the chokepoint (`charts.try_chart_png`): a shape
            refusal, leak finding, or missing rasterizer on this machine
            drops the chart silently and the section renders byte-identical
            to pre-OUT3 — which is why `charts` NEVER satisfies the
            section-content requirement below: every charted section must
            also carry its table/tile/body representation of the same
            numbers (the fallback is structural).
        A section may mix body / bullets / table / matrix / tiles / timeline.
        Render order is: tiles → charts → body → bullets → table → matrix →
        timeline.
      footer_text: centered footer text on every page. None (the default) uses
        the brand's `footer_line` ("Command Room" on the default brand). Pass a
        string to override. NEVER include provenance metadata
        (Source / Fired / meeting_id / TTL) per CONTRACT.md Rule 15.
      contract: output-contract gate mode (SPEC B3). Runs BEFORE the voice
        gate (canonical order: contract → voice → render → leak scan).
          "enforce" — raise OutputContractError (NO file written) on any
            BLOCKING contract violation (missing required section, too-few
            bullets, blank matrix/KPI cell, placeholder text, over-cap word
            count). Warn-severity violations (e.g. the call_prep total-word
            floor) print to stderr and the save proceeds — a sparse client
            workspace is never blocked for lack of substrate.
          "report"  — never block; print every violation to stderr and save.
            For backfill/regeneration and orchestrators mid-incident.
          "off"     — skip the contract gate entirely.
      contract_profile: optional alternate rule set passed through to the
        validator (e.g. "call_prep_internal" for internal-only meetings,
        which relaxes the total-word floor and uses the internal section list).
      voice_gate: pre-save voice-tell gate mode (SPEC B2).
          "default" — hard-fail (raise VoiceTellError, NO file written) when a
            fail-severity banned phrase appears AND brief_kind is a canonical
            outbound kind (memo / one_pager / decision_memo / board_pack /
            followup_pack). All other kinds are warn-only.
          "warn"    — never block; emit findings to stderr and save anyway.
            Escape valve for a caller with a legitimate calibrated override.
          "off"     — skip the voice gate entirely.
      exec_header: dict (SPEC EXEC1 element 1 — the 30-second contract).
        Keys: `verdict` (one bold conclusion sentence), `changed`, `decide`,
        `needs`. Rendered as the first block before any section. REQUIRED
        (with a non-empty verdict) for STANDARD_KINDS since the OUT2 §4 flip —
        omitting it raises ValueError BEFORE Document() is built (a brief_meta
        severity="error" audit event records the refusal). Other kinds may
        pass it freely; it's never required for them.
      asks: optional list of {text, deadline?} dicts (SPEC EXEC1 element 4 —
        the ASK block). Max MAX_ASKS (3); more raises ValueError. Each is a
        reader-action; rendered last under "What I need from you". Empty / None
        → nothing rendered (the exec header already said NEEDED Nothing). When a
        widget is the action surface, DON'T also pass asks — one-ask-surface.
      workspace_root: optional absolute workspace root. When provided, the
        exec-standard audit ALSO appends a brief_meta event to events.jsonl
        (best-effort, never blocks or masks the outcome). Omit it and the
        audit is stderr-only — still never user-visible. ALSO the
        source of the resolved brand theme when `brand` is not passed (reads
        workspace.brand / orgs[org_id].brand from entities.json — SPEC OUT1).
      brand: optional resolved brand dict (SPEC OUT1) overriding the theme for
        THIS render, highest precedence. Normally omitted — the theme resolves
        from workspace_root's entities.json (byte-stable DEFAULT_BRAND when no
        brand object is configured). Pass an explicit dict for a per-document
        theme (e.g. a per-org rendering assembled by the caller).
      org_id: optional org id whose per-org `brand` overrides the workspace
        brand, for a document scoped to one client org. Only consulted when
        `brand` is not passed and `workspace_root` is set.

    Returns: `output_path` on success.

    Raises:
      ValueError on bad inputs.
      OutputContractError (SPEC B3) when the pre-save contract gate blocks —
        raised BEFORE Document() is built, so no partial file exists and no
        content is lost; the caller reads each violation's section + fix_hint,
        rewrites ONLY the failing sections, and re-saves.
      VoiceTellError (SPEC B2) when the pre-save voice gate blocks — raised
        BEFORE Document() is built, so no partial file exists and no content
        is lost; the caller rewrites the flagged lines and re-saves.
      Lets python-docx errors propagate.

    Canonical pre-save gate order (B2 / B3): input validation → contract gate →
    voice gate → render (Document) → post-render leak scan.
    """
    # SPEC OUT5 §3b — the canonical pre-save gate sequence (input validation →
    # EXEC1 kwarg validation → rec-ordering → contract gate → voice gate →
    # exec-header requirement) runs through the SHARED stack in brief_gates.py,
    # the same call the premium-HTML backend makes. Every raise happens before
    # Document() is built — no partial file, no content lost. Do NOT add a gate
    # inline here: it belongs in brief_gates.run_pre_save_gates, or the G16
    # parity guard fails naming the backend that lags.
    eyebrow_by_kind = EYEBROW_BY_KIND
    gates_ran: List[str] = _run_pre_save_gates(
        brief_kind=brief_kind,
        title=title,
        subtitle=subtitle,
        sections=sections,
        supported_kinds=SUPPORTED_BRIEF_KINDS,
        contract=contract,
        contract_profile=contract_profile,
        voice_gate=voice_gate,
        exec_header=exec_header,
        asks=asks,
        workspace_root=workspace_root,
    )

    # SPEC OUT1 — resolve + apply the render theme (brand) for THIS render.
    # Precedence: explicit `brand=` dict > the workspace's resolved brand
    # (workspace.brand + optional orgs[org_id].brand from entities.json) >
    # byte-stable DEFAULT_BRAND. `get_brand(None)` is pure defaults with no I/O,
    # so a zero-config workspace renders the (upgraded) default theme. The
    # render-theme globals are restored in the finally so a later render in the
    # same process is never contaminated (defaults-are-defaults invariant).
    resolved_brand = brand if brand is not None else get_brand(workspace_root, org_id)
    _apply_brand(resolved_brand)

    # SPEC OUT2 §5 — resolve + apply the cross-skill output profile for THIS
    # render. Absent / unconfigured profile == DEFAULT_OUTPUT_PROFILE == today's
    # behavior, byte-stably (same invariant as the brand). Restored in the
    # finally alongside the brand globals.
    resolved_profile = get_output_profile(workspace_root)
    _apply_output_profile(resolved_profile)

    # page_cap (WARN-ONLY, forever): shared with the premium-HTML backend so a
    # configured cap warns identically whichever backend renders (SPEC OUT5).
    _warn_page_cap(resolved_profile, brief_kind, title, subtitle, sections)

    try:
        doc = Document()
        _set_normal_baseline(doc)
        _tighten_margins(doc)

        # SPEC OUT1 §3b — brand letterhead. Renders ONLY when logo_path is set
        # AND the file exists (R26: a missing logo silently falls back to the
        # quiet no-logo header, never an error). Default brand has no logo, so
        # this is a no-op on a fresh workspace — the header is byte-unchanged.
        logo_abspath = _resolve_logo(resolved_brand.get("logo_path"), workspace_root)
        if logo_abspath:
            _add_logo_header(doc, logo_abspath)

        eyebrow_label = eyebrow_by_kind[brief_kind]
        _add_eyebrow(doc, eyebrow_label)
        _add_title(doc, title)
        _add_subtitle(doc, subtitle)
        _add_header_rule(doc)

        # SPEC EXEC1 element 1 — the 30-second contract renders before any section.
        # FS-13: document/decision kinds get the verdict lead only, no eyebrow.
        if exec_header and (exec_header.get("verdict") or "").strip():
            _add_exec_header(doc, exec_header, brief_kind=brief_kind)

        for sec in sections:
            heading = sec.get("heading")
            if not heading:
                raise ValueError(f"section missing 'heading': {sec!r}")
            _add_section_heading(doc, heading)
            body = sec.get("body")
            bullets = sec.get("bullets")
            table = sec.get("table")
            matrix = sec.get("matrix")
            tiles = sec.get("tiles")
            timeline = sec.get("timeline")
            charts_spec = sec.get("charts")
            if tiles and not isinstance(tiles, list):
                raise ValueError(f"section 'tiles' must be a list: {sec!r}")
            if body and not isinstance(body, str):
                raise ValueError(f"section 'body' must be a string: {sec!r}")
            if charts_spec and not isinstance(charts_spec, list):
                raise ValueError(f"section 'charts' must be a list: {sec!r}")
            # SPEC OUT2 §5 — visual_bias sets the tiles/body order within a
            # section. Default "tiles_first" is the pre-profile order
            # (tiles -> body), byte-stably; "prose_first" flips just these two
            # (bullets/table/matrix/timeline order is unchanged). SPEC OUT3:
            # charts ride directly under the tile band in both orders.
            if _VISUAL_BIAS == "prose_first":
                if body:
                    _add_body_paragraphs(doc, body)
                if tiles:
                    _add_stat_tiles(doc, tiles)
                if charts_spec:
                    _add_charts(doc, charts_spec, resolved_brand)
            else:
                if tiles:
                    _add_stat_tiles(doc, tiles)
                if charts_spec:
                    _add_charts(doc, charts_spec, resolved_brand)
                if body:
                    _add_body_paragraphs(doc, body)
            if bullets:
                if not isinstance(bullets, list):
                    raise ValueError(f"section 'bullets' must be a list: {sec!r}")
                _add_bullets(doc, bullets)
            if table:
                if not isinstance(table, dict):
                    raise ValueError(f"section 'table' must be a dict: {sec!r}")
                _add_table(
                    doc,
                    rows=table["rows"],
                    headers=table.get("headers"),
                    highlight_row_idx=table.get("highlight_row_idx"),
                    column_widths=table.get("column_widths"),
                )
            if matrix:
                if not isinstance(matrix, dict):
                    raise ValueError(f"section 'matrix' must be a dict: {sec!r}")
                _add_matrix(
                    doc,
                    cells=matrix["cells"],
                    headers_row=matrix.get("headers_row"),
                    headers_col=matrix.get("headers_col"),
                    star_col_idx=matrix.get("star_col_idx"),
                    flag_col_idx=matrix.get("flag_col_idx"),
                )
            if timeline:
                if not isinstance(timeline, list):
                    raise ValueError(f"section 'timeline' must be a list: {sec!r}")
                _add_timeline(doc, timeline)
            if not body and not bullets and not table and not matrix \
                    and not tiles and not timeline:
                # SPEC OUT3: 'charts' deliberately does NOT satisfy this —
                # charts render best-effort, so a chart-only section would be
                # an empty frame on any machine without a rasterizer. Every
                # charted section carries its fallback representation.
                raise ValueError(
                    f"section needs 'body', 'bullets', 'table', 'matrix', "
                    f"'tiles', or 'timeline' ('charts' is best-effort and "
                    f"never stands alone): {sec!r}"
                )

        # SPEC EXEC1 element 4 — the ASK block renders last, after all sections.
        # Zero asks → nothing (the exec header already said NEEDED Nothing).
        if asks:
            renderable = [a for a in asks if (a.get("text") or "").strip()]
            if renderable:
                _add_asks_block(doc, renderable)

        # Footer: caller override wins; otherwise the brand footer_line (which
        # is "Command Room" on the default brand).
        _add_footer(doc, footer_text if footer_text is not None else FOOTER_DEFAULT)
        doc.save(output_path)

        # v3.13.8+ — universal post-render leak scan gate (Bug #57 + #59 + #54).
        # Runs against every brief regardless of which skill called us. The
        # scanner is added by §2.4. We invoke it lazily to avoid a hard
        # dependency cycle during partial-install scenarios.
        try:
            from docx_leak_scanner import scan_docx_for_leaks, LeakScanError
            scan_docx_for_leaks(output_path)
            gates_ran.append("leak")
        except ImportError:
            # docx_leak_scanner not installed yet (e.g. a workspace that hasn't
            # taken the v3.13.8 update). Don't block briefs; the scanner will
            # apply on the next plugin update.
            pass

        # SPEC GATE1 — emit the detectable-bypass audit AFTER a fully successful
        # render+save (deliverable on disk, all wired gates passed). A composer
        # fire that produces a doc with NO gate_ran event for that turn is a
        # flaggable bypass. Best-effort + never raises (deliverable already valid).
        _emit_gate_ran_audit(brief_kind, gates_ran, output_path, workspace_root)

        return output_path
    finally:
        # Restore defaults so the next render in this process is never
        # contaminated by this render's theme or profile (byte-stable defaults
        # invariant).
        _apply_brand(_DEFAULT_RESOLVED)
        _apply_output_profile(DEFAULT_OUTPUT_PROFILE)


def make_brief_from_json(json_payload: str) -> str:
    """JSON wrapper for orchestrator bash invocations.

    Pipe a JSON object on stdin OR pass as the first CLI arg.
    Required keys mirror `make_brief()` kwargs.

    Returns the output path (also printed to stdout for shell capture).
    """
    payload = json.loads(json_payload)
    path = make_brief(
        payload["output_path"],
        brief_kind=payload["brief_kind"],
        title=payload["title"],
        subtitle=payload["subtitle"],
        sections=payload["sections"],
        footer_text=payload.get("footer_text"),
        voice_gate=payload.get("voice_gate", "default"),
        contract=payload.get("contract", "enforce"),
        contract_profile=payload.get("contract_profile"),
        exec_header=payload.get("exec_header"),
        asks=payload.get("asks"),
        workspace_root=payload.get("workspace_root"),
        brand=payload.get("brand"),
        org_id=payload.get("org_id"),
    )
    print(path)
    return path


__all__ = ["make_brief", "make_brief_from_json"]


if __name__ == "__main__":
    # CLI: `python3 brief_writer.py '<json>'` OR pipe JSON on stdin.
    if len(sys.argv) > 1:
        make_brief_from_json(sys.argv[1])
    else:
        make_brief_from_json(sys.stdin.read())
